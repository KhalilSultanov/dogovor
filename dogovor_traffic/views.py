import os
from django.http import HttpResponse
from django.shortcuts import render
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, Inches


def replace_paragraph_text_with_styles(paragraph, new_text):
    """
    Заменяет текст в абзаце, сохраняя исходное форматирование.
    """
    if paragraph.runs:
        style = paragraph.runs[0].style
        font = paragraph.runs[0].font
        is_bold = font.bold
        font_name = font.name
        font_size = font.size
    else:
        is_bold = False
        font_name = 'Calibri'
        font_size = Pt(9)

    paragraph.clear()
    run = paragraph.add_run(new_text)
    run.bold = is_bold
    run.font.name = font_name
    run.font.size = font_size




def replace_analytics_tags(doc, analitic_system, analitic_system_user, system_search):
    # Замена для "Яндекс.Вебмастер" и "Search Console"
    if 'yandex_web' in analitic_system:
        yandex_web_text = '«Яндекс.Вебмастер»'
    else:
        yandex_web_text = ''

    if 'search_console' in analitic_system:
        search_console_text = '«Search Console»'
    else:
        search_console_text = ''

    # Если выбраны обе системы, соединяем их с " и "
    if yandex_web_text and search_console_text:
        web_analytics = f'{yandex_web_text} и {search_console_text}'
    else:
        web_analytics = yandex_web_text or search_console_text
    replace_tag_with_text(doc, '{YANDEX_WEB}', web_analytics)
    replace_tag_with_text(doc, '{SEARCH_CONSOLE}', '')  # Тег удаляем, текст уже вставлен

    # Замена для "Яндекс.Метрика" и "Google Analytics"
    yandex_metric_text = 'Яндекс.Метрика' if 'yandex_metric' in analitic_system_user else ''
    google_analytic_text = 'Google Analytics' if 'google_analitic' in analitic_system_user else ''
    user_analytics = ', '.join(filter(None, [yandex_metric_text, google_analytic_text]))
    replace_tag_with_text(doc, '{YANDEX_METRIC}', user_analytics)
    replace_tag_with_text(doc, '{GOOGLE_ANALITIC}', '')  # Тег удаляем, текст уже вставлен

    # Замена для "Яндекс" и "Google" в контексте поисковой системы
    yandex_search_text = 'Яндекс.Метрика' if 'yandex_system' in system_search else ''
    google_search_text = 'Google Analytics' if 'google_system' in system_search else ''
    search_analytics = ' и '.join(filter(None, [yandex_search_text, google_search_text]))
    replace_tag_with_text(doc, '{YANDEX}', search_analytics)
    replace_tag_with_text(doc, '{GOOGLE}', '')  # Тег удаляем, текст уже вставлен

def handle_additional_work_sections(doc, platform_choice):

    for paragraph in doc.paragraphs:
        if '{WORD_PRESS}' in paragraph.text or '{NOT_WORD_PRESS}' in paragraph.text:
            if platform_choice == 'wordpress':
                text = paragraph.text.replace('{WORD_PRESS}', 'Работы программиста по результатам коммерческого аудита. Работы программиста после проведения других аудитов включены в счёт.').replace('{NOT_WORD_PRESS}', '')
            elif platform_choice == 'not_wordpress':
                text = paragraph.text.replace('{WORD_PRESS}', '').replace('{NOT_WORD_PRESS}', 'Работы программиста.')
            else:
                text = paragraph.text.replace('{WORD_PRESS}', '').replace('{NOT_WORD_PRESS}', '')
            replace_paragraph_text_with_styles(paragraph, text)


def handle_conditional_sections(doc, edo):
    edo_text_1 = "(в том числе его получения с использованием системы электронного документооборота)" if edo == "YES" else ""
    edo_text_2 = (
                "\n9.4. Стороны согласовали, что они вправе осуществлять документооборот в электронном виде по телекоммуникационным каналам связи с использованием усиленной квалификационной электронной подписи посредством системы электронного документооборота СБИС. " + "\n" +
                "9.4.1. В целях настоящего договора под электронным документом понимается документ, созданный в электронной форме без предварительного документирования на бумажном носителе, подписанный электронной подписью в порядке, установленном законодательством Российской Федерации. Стороны признают электронные документы, заверенные электронной подписью, при соблюдении требований Федерального закона от 06.04.2011 № 63-ФЗ 'Об электронной подписи' юридически эквивалентными документам на бумажных носителях, заверенным соответствующими подписями и оттиском печатей Сторон." + "\n" +
                "9.5. Все изменения и дополнения к договору оформляются в виде дополнений и приложений к договору, являющийся его неотъемлемой частью." + "\n" +
                "9.6. Договор составлен в двух подлинных экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из сторон. ") \
        if edo == "YES" else ""
    not_edo_text = "на почту Исполнителя" if edo == "NO" else ""

    write_by_hand = (
                "\n9.4. Все изменения и дополнения к договору оформляются в виде дополнений и приложений к договору, являющийся его неотъемлемой частью. " + '\n' +
                "9.5. Договор составлен в двух подлинных экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из сторон.)") if edo == "NO" else ""

    replacements = {
        '{EDO_1}': edo_text_1,
        '{EDO_2}': edo_text_2,
        '{NOT_EDO}': not_edo_text,
        '{WRITE_BY_HAND}': write_by_hand
    }


    for paragraph in doc.paragraphs:
        for tag, replacement in replacements.items():
            if tag in paragraph.text:
                paragraph_text = paragraph.text.replace(tag, replacement)
                replace_paragraph_text_with_styles(paragraph, paragraph_text)

def replace_tag_with_text(doc, tag, text=None):
    for paragraph in doc.paragraphs:
        if tag in paragraph.text:
            new_text = paragraph.text.replace(tag, text if text is not None else '').strip()
            if new_text:
                replace_paragraph_text_with_styles(paragraph, new_text)
            else:
                # Если параграф остался пустым, удаляем его
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None

def replace_underscores_with_signature(doc, placeholder_text, signature_path):
    def replace_in_paragraph(paragraph):
        runs = paragraph.runs
        full_text = ''.join(run.text for run in runs)
        if placeholder_text in full_text:
            split_text = full_text.split(placeholder_text)
            # Clear existing runs
            for run in runs:
                run.text = ""
            # Add new runs with the signature and styled text
            paragraph.add_run(split_text[0])
            paragraph.add_run().add_picture(signature_path, width=Inches(0.8))
            run = paragraph.add_run("Михайлов Д.С.")
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            if len(split_text) > 1:
                paragraph.add_run(split_text[1])

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            replace_in_paragraph(paragraph)
        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_in_paragraph(paragraph)


def find_and_offset_director_text(doc):
    # Обход всех параграфов
    for paragraph in doc.paragraphs:
        if "________________{FIO_DIRECTOR}" in paragraph.text:
            paragraph.text = "\n\n\n" + paragraph.text

    # Обход всех секций (header и footer)
    for section in doc.sections:
        for header in section.header.paragraphs:
            if "________________{FIO_DIRECTOR}" in header.text:
                header.text = "\n\n\n" + header.text
        for footer in section.footer.paragraphs:
            if "________________{FIO_DIRECTOR}" in footer.text:
                footer.text = "\n\n\n" + footer.text

    # Обход всех таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "________________{FIO_DIRECTOR}" in paragraph.text:
                        paragraph.text = "\n\n\n" + paragraph.text

def add_newline_before_text(doc, search_text):
    """
    Вставляет новую строку перед указанным текстом в документе.
    """
    for paragraph in doc.paragraphs:
        if search_text in paragraph.text:
            before_text, after_text = paragraph.text.split(search_text, 1)
            new_text = before_text.rstrip() + '\n' + search_text + after_text
            replace_paragraph_text_with_styles(paragraph, new_text)

def handle_search_engine(doc, search_engine_choice):
    if search_engine_choice == 'yandex':
        replace_tag_with_text(doc, '{YANDEX_1}', '4.3. Отчеты о посещениях Интернет-сайта составляются на основании данных Яндекс.Метрика по показателю «Посетители». ')
        replace_tag_with_text(doc, '{GOOGLE_1}', '')
        replace_tag_with_text(doc, '{YANDEX_2}', '5.10. В случае если в период выполнения работ имела место неработоспособность системы анализа Яндекс.Метрика, установленная на Интернет-сайте, по вине любой из сторон, то расчет количества посещений осуществляется Исполнителем согласно данных других аналитических систем, о чем делается соответствующая запись в направляемом Заказчику отчете.')
        replace_tag_with_text(doc, '{GOOGLE_2}', '')
        replace_tag_with_text(doc, '{YANDEX}', 'Яндекс.Метрика')
        replace_tag_with_text(doc, '{GOOGLE}', '')
    elif search_engine_choice == 'google':
        replace_tag_with_text(doc, '{GOOGLE_1}', '4.3. Отчеты о посещениях Интернет-сайта составляются на основании данных Google Analytics по показателю «Сеансы». ')
        replace_tag_with_text(doc, '{YANDEX_1}', '')
        replace_tag_with_text(doc, '{YANDEX_2}', '')
        replace_tag_with_text(doc, '{GOOGLE_2}', '5.10. В случае если в период выполнения работ имела место неработоспособность системы анализа Google Analytics, установленная на Интернет-сайте, по вине любой из сторон, то расчет количества посещений осуществляется Исполнителем согласно данных других аналитических систем, о чем делается соответствующая запись в направляемом Заказчику отчете.')
        replace_tag_with_text(doc, '{GOOGLE}', 'Google Analytics')
        replace_tag_with_text(doc, '{YANDEX}', '')
def process_contract(request):
    if request.method == 'POST':


        service_to_tag_mapping = {

            'optimization_headers': ('{HEAD_PAGES_1}', '- Оптимизация заголовков страниц;'),
            'optimization_headers2': (
            '{HEAD_PAGES_2}', '- Оптимизация заголовков страниц и техническое задание на их внедрение;'),
            'optimization_headers3': ('{HEAD_PAGES_3}', ''),

            'optimization_metatags': ('{METATAGS_1}', '- Оптимизация метатегов;'),
            'optimization_metatags2': (
            '{METATAGS_2}', '- Оптимизация метатегов и техническое задание на их внедрение;'),
            'optimization_metatags3': ('{METATAGS_3}', ''),

            'writing_optimization': ('{NEURO_1}', '- Написание текстов с помощью нейросетей и их оптимизация;'),
            'writing_optimization2': ('{NEURO_2}', '- Техническое задание на написание текстов и их оптимизация;'),
            'writing_optimization3': ('{NEURO_3}', ''),

            'site_structure_optimization': ('{STRUCTURE_1}', '- Оптимизация структуры сайта;'),
            'site_structure_optimization2': (
            '{STRUCTURE_2}', '- Оптимизация структуры сайта и техническое задание на внедрение;'),
            'site_structure_optimization3': ('{STRUCTURE_3}', ''),

            'technical_error_fixing': ('{FIX_1}', '- Устранение технических ошибок на сайте;'),
            'technical_error_fixing2': ('{FIX_2}', '- Техническое задание на устранение технических ошибок на сайте;'),
            'technical_error_fixing3': ('{FIX_3}', ''),

            'design_layouts': (
            '{TZ_1}', '- Техническое задание (ТЗ) на создание дизайн-макетов отдельных блоков или страниц на сайте;'),

            'creating_pages': ('{CREATE_PAGES_1}', '- Создание страниц на сайте;'),
            'creating_pages2': ('{CREATE_PAGES_2}', '- Техническое задание на создание страниц на сайте;'),
            'creating_pages3': ('{CREATE_PAGES_3}', '')
        }

        contract_number = request.POST.get('contract_number')
        date_day = request.POST.get('date_day')
        site_name = request.POST.get('site_name')

        price_count_digit = request.POST.get('price_count_digit')
        price_count_word = request.POST.get('price_count_word')
        search_engine_choice = request.POST.get('search_engine', None)

        analitic_system = request.POST.getlist('analitic_system')
        analitic_system_user = request.POST.getlist('analitic_system_user')

        system_search = request.POST.getlist('search_system')

        pay_for_site = request.POST.get('pay_for_site')
        prime = request.POST.get('prime')
        visit_count = request.POST.get('visit_count')
        start_date = request.POST.get('start_date')
        end_date = request.POST.get('end_date')
        red_organization_name = request.POST.get('red_organization_name')


        date_month = request.POST.get('date_month')
        date_year = request.POST.get('date_year')
        organization_name = request.POST.get('organization_name')
        reason = request.POST.get('reason')
        person_name = request.POST.get('person_name')
        director_name = request.POST.get('director_name')
        email = request.POST.get('email')
        inn = request.POST.get('inn')
        ogrn = request.POST.get('ogrn')
        registration_address = request.POST.get('registration_address')
        checking_account = request.POST.get('checking_account')
        correspondent_account = request.POST.get('correspondent_account')
        bank_name = request.POST.get('bank_name')
        bic = request.POST.get('bic')
        edo = request.POST.get('edo')



        support_options = request.POST.getlist('support[]')
        platform_choice = request.POST.get('platform', None)
        selected_services = request.POST.getlist('services[]')


        template_filename = 'Договор Трафик метки.docx'
        template_path = os.path.join(os.path.dirname(__file__), '../dogovora', template_filename)
        doc = Document(template_path)
        signature_image_path = os.path.join(os.path.dirname(__file__), '../dogovora/podpis.jpg')

        if 'support_site' in support_options:
            support_text = '- Поддержка сайта в техническом плане;'
            replace_tag_with_text(doc, '{SITE_SUPPORT}', support_text)
        else:
            replace_tag_with_text(doc, '{SITE_SUPPORT}')

        used_tags = set()

        for service_key in selected_services:
            if service_key in service_to_tag_mapping:
                tag, text = service_to_tag_mapping[service_key]
                replace_tag_with_text(doc, tag, text)
                used_tags.add(tag)

        all_tags = set(service_to_tag_mapping.values())
        unused_tags = {tag for tag, _ in all_tags if tag not in used_tags}
        for tag in unused_tags:
            replace_tag_with_text(doc, tag)

        handle_conditional_sections(doc, edo)
        handle_additional_work_sections(doc, platform_choice)
        handle_search_engine(doc, search_engine_choice)
        replace_analytics_tags(doc, analitic_system, analitic_system_user, system_search)


        footer = doc.sections[0].footer
        for paragraph in footer.paragraphs:
            paragraph.alignment = 1
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
        paragraph = footer.paragraphs[0]
        paragraph.add_run("________________" + director_name + "").font.size = Pt(12)
        paragraph.add_run("                                                                           ").font.size = Pt(
            12)

        if edo == "YES":
            paragraph.add_run("________________Михайлов Д.С.").font.size = Pt(12)
        else:
            replace_underscores_with_signature(doc, "________________Михайлов Д.С.", signature_image_path)
            find_and_offset_director_text(doc)
            run = paragraph.add_run()
            run.add_picture(signature_image_path, width=Inches(0.8))  # Настройте ширину по необходимости
            paragraph.add_run("Михайлов Д.С.").font.size = Pt(12)

        add_newline_before_text(doc, "10.5. Договор составлен")
        replacements = {
            '{DOGOVOR_NUMBER}': contract_number,
            '{DAY}': date_day,
            '{MONTH}': date_month,
            '{YEAR}': date_year,
            '{CUSTOMER_ORGANIZATION}': organization_name,
            '{CUSTOMER_FIO}': person_name,
            '{DOGOVOR_OSNOVANIE}': reason,

            '{SITE_NAME}': site_name,
            '{PRICE_COUNT}': price_count_digit,
            '{PRICE_COUNT_IN_WORDS}': price_count_word,
            '{PAY_FOR_SITE}': pay_for_site,
            '{PAY_FOR_ONCE}': prime,
            '{VISIT_COUNT}': visit_count,
            '{START}': start_date,
            '{END}': end_date,

            '{FIO_DIRECTOR}': director_name,
            '{CUSTOMER_EMAIL}': email,
            '{CUSTOMER_NAME}': organization_name,
            '{RED_CUSTOMER_ORGANIZATION}': red_organization_name,

            '{INN}': inn,
            '{OGRN}': ogrn,
            '{REGISTRATION_ADDRESS}': registration_address,
            '{PAYMENT_ACCOUNT}': checking_account,
            '{CORRESPONDENT}': correspondent_account,
            '{BANK_NAME}': bank_name,
            '{BIK}': bic,
        }

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)
                        if key == '{DOGOVOR_NUMBER}' or key == '{CUSTOMER_NAME}':
                            run.bold = True


        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)
                                for run in paragraph.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(9)
                                    if run.text.strip() == value:
                                        run.bold = True


        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9)


        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="processed_contract.docx"'
        doc.save(response)
        return response
    else:
        return render(request, 'dogovor_traffic_form.html')
