import os
from django.http import HttpResponse
from django.shortcuts import render
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, Inches
from num2words import num2words

from docx.shared import Pt

from docx.shared import Pt

from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


def insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_paragraph.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
    if style:
        new_paragraph.style = style
    return new_paragraph


def add_guarantee_section(doc):
    # Флаг, который указывает, что мы находимся внутри раздела 4
    in_section_4 = False

    # Текст новых пунктов
    new_section_header = '4. Ключевые фразы и гарантия'
    new_point_4_4 = '4.4. Гарантией выполненных работ является вывод в топ-10 поисковой сети Яндекс не менее 30% запросов, указанных в Приложении №2, в течение 6 (шести) месяцев работ.'
    new_point_4_5 = '4.5. При невыполнении гарантийных обязательств, Исполнитель предоставляет 7 (седьмой) месяц работ бесплатно.'
    new_point_4_6 = '''4.6. Условия гарантии являются недействительными в случаях:
- действий или бездействий поисковых систем, хостинга, системы сайта – для этого важно настраивать журнал логов;
- форс-мажорных обстоятельств, указанных в п. 8 настоящего Договора;
- работ Заказчика и третьих лиц на сайте без согласования с Исполнителем;
- смены стратегии по инициативе Заказчика;
- несоблюдения/несвоевременного соблюдения выполнения Заказчиком данных Исполнителем рекомендаций по п.1.4. настоящего Договора.'''

    paragraphs = doc.paragraphs
    i = 0
    while i < len(paragraphs):
        paragraph = paragraphs[i]
        text = paragraph.text.strip()

        if text.startswith('4. '):
            # Находим заголовок раздела 4 и заменяем его
            paragraph.clear()
            run = paragraph.add_run(new_section_header)
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            run.bold = True
            in_section_4 = True
            i += 1
            continue

        if in_section_4:
            if text.startswith('5. '):
                # Дошли до раздела 5, выходим из раздела 4
                in_section_4 = False
                i += 1
                continue

            if text.startswith('4.4. '):
                # Заменяем пункт 4.4 и добавляем новые пункты
                paragraph.clear()
                run = paragraph.add_run(new_point_4_4)
                run.font.name = 'Calibri'
                run.font.size = Pt(9)

                # Вставляем новый параграф для 4.5
                p4_5 = insert_paragraph_after(paragraph, new_point_4_5)
                # Вставляем параграфы для 4.6 и его подпунктов
                lines = new_point_4_6.split('\n')
                previous_p = p4_5
                for line in lines:
                    p = insert_paragraph_after(previous_p, line)
                    previous_p = p  # Обновляем ссылку на предыдущий параграф

                # После добавления новых пунктов, пропускаем оригинальный 4.5 и 4.6, если они есть
                # Проверяем следующие параграфы
                while i + 1 < len(paragraphs) and paragraphs[i + 1].text.strip().startswith('4.'):
                    i += 1
                i += 1
                continue

            else:
                i += 1
                continue
        else:
            i += 1


def replace_text_in_paragraphs(doc, search_text, replace_text):
    for paragraph in doc.paragraphs:
        if search_text in paragraph.text:
            new_text = paragraph.text.replace(search_text, replace_text)
            replace_paragraph_text_with_styles(paragraph, new_text)


def remove_bullet_point(doc, bullet_text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == bullet_text:
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None
            break  # Предполагаем, что такой пункт единственный


def make_text_bold_in_doc(doc, search_text):
    main_text = search_text.split(", именуемое")[0]  # Основная часть текста до ", именуемое"

    for paragraph in doc.paragraphs:
        if main_text in paragraph.text:
            runs = paragraph.runs
            for run in runs:
                if main_text in run.text:
                    split_text = run.text.split(main_text, 1)
                    run.text = split_text[0]
                    bold_run = paragraph.add_run(main_text)
                    bold_run.bold = True
                    bold_run.font.name = 'Calibri'
                    bold_run.font.size = Pt(9)
                    if len(split_text) > 1:
                        after_bold_run = paragraph.add_run(split_text[1])
                        after_bold_run.font.name = 'Calibri'
                        after_bold_run.font.size = Pt(9)

    # Повторение для таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if main_text in paragraph.text:
                        runs = paragraph.runs
                        for run in runs:
                            if main_text in run.text:
                                split_text = run.text.split(main_text, 1)
                                run.text = split_text[0]
                                bold_run = paragraph.add_run(main_text)
                                bold_run.bold = True
                                bold_run.font.name = 'Calibri'
                                bold_run.font.size = Pt(9)
                                if len(split_text) > 1:
                                    after_bold_run = paragraph.add_run(split_text[1])
                                    after_bold_run.font.name = 'Calibri'
                                    after_bold_run.font.size = Pt(9)


def replace_paragraph_text_with_styles(paragraph, new_text):
    """
    Заменяет текст в абзаце, сохраняя исходное форматирование.
    """
    if paragraph.runs:
        style = paragraph.runs[0].style
        font = paragraph.runs[0].font
        is_bold = font.bold
        font_name = font.name
        font_size = font.size
    else:
        is_bold = False
        font_name = 'Calibri'
        font_size = Pt(9)

    paragraph.clear()
    run = paragraph.add_run(new_text)
    run.bold = is_bold
    run.font.name = font_name
    run.font.size = font_size


def replace_analytics_tags(doc, analitic_system, analitic_system_user, system_search):
    # Замена для "Яндекс.Вебмастер" и "Search Console"
    if 'yandex_web' in analitic_system:
        yandex_web_text = '«Яндекс.Вебмастер»'
    else:
        yandex_web_text = ''

    if 'search_console' in analitic_system:
        search_console_text = '«Search Console»'
    else:
        search_console_text = ''

    # Если выбраны обе системы, соединяем их с " и "
    if yandex_web_text and search_console_text:
        web_analytics = f'{yandex_web_text} и {search_console_text}'
    else:
        web_analytics = yandex_web_text or search_console_text
    replace_tag_with_text(doc, '{YANDEX_WEB}', web_analytics)
    replace_tag_with_text(doc, '{SEARCH_CONSOLE}', '')  # Тег удаляем, текст уже вставлен

    # Замена для "Яндекс.Метрика" и "Google Analytics"
    yandex_metric_text = 'Яндекс.Метрика' if 'yandex_metric' in analitic_system_user else ''
    google_analytic_text = 'Google Analytics' if 'google_analitic' in analitic_system_user else ''
    user_analytics = ', '.join(filter(None, [yandex_metric_text, google_analytic_text]))
    replace_tag_with_text(doc, '{YANDEX_METRIC}', user_analytics)
    replace_tag_with_text(doc, '{GOOGLE_ANALITIC}', '')  # Тег удаляем, текст уже вставлен

    # Замена для "Яндекс" и "Google" в контексте поисковой системы
    yandex_search_text = 'Яндекс' if 'yandex_system' in system_search else ''
    google_search_text = 'Google' if 'google_system' in system_search else ''
    search_analytics = ' и '.join(filter(None, [yandex_search_text, google_search_text]))
    replace_tag_with_text(doc, '{YANDEX}', search_analytics)
    replace_tag_with_text(doc, '{GOOGLE}', '')  # Тег удаляем, текст уже вставлен


def replace_tag_with_text(doc, tag, text):
    for paragraph in doc.paragraphs:
        if tag in paragraph.text:
            new_text = paragraph.text.replace(tag, text)
            replace_paragraph_text_with_styles(paragraph, new_text)
            if not new_text.strip():  # Если параграф остался пустым, удаляем его
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None


def handle_additional_work_sections(doc, platform_choice):
    for paragraph in doc.paragraphs:
        if '{WORD_PRESS}' in paragraph.text or '{NOT_WORD_PRESS}' in paragraph.text:
            if platform_choice == 'wordpress':
                text = paragraph.text.replace('{WORD_PRESS}',
                                              ' по результатам коммерческого аудита. Работы программиста после проведения других аудитов включены в счёт.').replace(
                    '{NOT_WORD_PRESS}', '')
            elif platform_choice == 'not_wordpress':
                text = paragraph.text.replace('{WORD_PRESS}', '').replace('{NOT_WORD_PRESS}',
                                                                          '.')
            else:
                text = paragraph.text.replace('{WORD_PRESS}', '').replace('{NOT_WORD_PRESS}', '')
            replace_paragraph_text_with_styles(paragraph, text)


def handle_conditional_sections(doc, edo):
    edo_text_1 = "(в том числе его получения с использованием системы электронного документооборота)" if edo == "YES" else ""
    edo_text_2 = (
            "\nЛибо посредством ЭДО: "
            "\n- ID Исполнителя в системе Тензор: 2BE894898d706174ab2aa3cdfc300550236"
            "\n- ID Заказчика: {CUSTOMER_ID} "
            "\n10.4. Стороны согласовали, что они вправе осуществлять документооборот в электронном виде по телекоммуникационным каналам связи с использованием усиленной квалификационной электронной подписи посредством системы электронного документооборота. " + "\n" +
            "10.4.1. В целях настоящего договора под электронным документом понимается документ, созданный в электронной форме без предварительного документирования на бумажном носителе, подписанный электронной подписью в порядке, установленном законодательством Российской Федерации. Стороны признают электронные документы, заверенные электронной подписью, при соблюдении требований Федерального закона от 06.04.2011 № 63-ФЗ 'Об электронной подписи' юридически эквивалентными документам на бумажных носителях, заверенным соответствующими подписями и оттиском печатей Сторон." + "\n" +
            "10.5. Все изменения и дополнения к договору оформляются в виде дополнений и приложений к договору, являющийся его неотъемлемой частью." + "\n" +
            "10.6. Договор составлен в двух подлинных экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из сторон, подписанных лично либо посредством ЭДО. ") \
        if edo == "YES" else ""
    not_edo_text = "на почту Исполнителя" if edo == "NO" else ""

    write_by_hand = (
            "\n10.4. Все изменения и дополнения к договору оформляются в виде дополнений и приложений к договору, являющийся его неотъемлемой частью. " + '\n' +
            "10.5. Договор составлен в двух подлинных экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из сторон.") if edo == "NO" else ""

    replacements = {
        '{EDO_1}': edo_text_1,
        '{EDO_2}': edo_text_2,
        '{NOT_EDO}': not_edo_text,
        '{WRITE_BY_HAND}': write_by_hand
    }

    for paragraph in doc.paragraphs:
        for tag, replacement in replacements.items():
            if tag in paragraph.text:
                paragraph_text = paragraph.text.replace(tag, replacement)
                replace_paragraph_text_with_styles(paragraph, paragraph_text)


def replace_tag_with_text(doc, tag, text=None):
    for paragraph in doc.paragraphs:
        if tag in paragraph.text:
            new_text = paragraph.text.replace(tag, text if text is not None else '').strip()
            if new_text:
                replace_paragraph_text_with_styles(paragraph, new_text)
            else:
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None


def replace_underscores_with_signature(doc, placeholder_text, signature_path):
    def replace_in_paragraph(paragraph):
        runs = paragraph.runs
        full_text = ''.join(run.text for run in runs)
        if placeholder_text in full_text:
            split_text = full_text.split(placeholder_text)
            for run in runs:
                run.text = ""
            paragraph.add_run(split_text[0])
            paragraph.add_run().add_picture(signature_path, width=Inches(0.8))
            run = paragraph.add_run("Михайлов Д.С.")
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            if len(split_text) > 1:
                paragraph.add_run(split_text[1])

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            replace_in_paragraph(paragraph)
        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_in_paragraph(paragraph)


def find_and_offset_director_text(doc):
    # Обход всех параграфов
    for paragraph in doc.paragraphs:
        if "________________{FIO_DIRECTOR}" in paragraph.text:
            paragraph.text = "\n\n\n" + paragraph.text

    # Обход всех секций (header и footer)
    for section in doc.sections:
        for header in section.header.paragraphs:
            if "________________{FIO_DIRECTOR}" in header.text:
                header.text = "\n\n\n" + header.text
        for footer in section.footer.paragraphs:
            if "________________{FIO_DIRECTOR}" in footer.text:
                footer.text = "\n\n\n" + footer.text

    # Обход всех таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "________________{FIO_DIRECTOR}" in paragraph.text:
                        paragraph.text = "\n\n\n" + paragraph.text


def add_newline_before_text(doc, search_text):
    """
    Вставляет новую строку перед указанным текстом в документе.
    """
    for paragraph in doc.paragraphs:
        if search_text in paragraph.text:
            before_text, after_text = paragraph.text.split(search_text, 1)
            new_text = before_text.rstrip() + '\n' + search_text + after_text
            replace_paragraph_text_with_styles(paragraph, new_text)


def process_contract(request):
    if request.method == 'POST':

        service_to_tag_mapping = {

            'optimization_headers': ('{HEAD_PAGES_1}', '- Оптимизация заголовков страниц;'),
            'optimization_headers2': (
                '{HEAD_PAGES_2}', '- Оптимизация заголовков страниц и техническое задание на их внедрение;'),
            'optimization_headers3': ('{HEAD_PAGES_3}', ''),

            'optimization_metatags': ('{METATAGS_1}', '- Оптимизация метатегов;'),
            'optimization_metatags2': (
                '{METATAGS_2}', '- Оптимизация метатегов и техническое задание на их внедрение;'),
            'optimization_metatags3': ('{METATAGS_3}', ''),

            'writing_optimization': ('{NEURO_1}', '- Написание текстов с помощью нейросетей и их оптимизация;'),
            'writing_optimization2': ('{NEURO_2}', '- Техническое задание на написание текстов и их оптимизация;'),
            'writing_optimization3': ('{NEURO_3}', ''),

            'site_structure_optimization': ('{STRUCTURE_1}', '- Оптимизация структуры сайта;'),
            'site_structure_optimization2': (
                '{STRUCTURE_2}', '- Оптимизация структуры сайта и техническое задание на внедрение;'),
            'site_structure_optimization3': ('{STRUCTURE_3}', ''),

            'technical_error_fixing': ('{FIX_1}', '- Устранение технических ошибок на сайте;'),
            'technical_error_fixing2': ('{FIX_2}', '- Техническое задание на устранение технических ошибок на сайте;'),
            'technical_error_fixing3': ('{FIX_3}', ''),

            'design_layouts': (
                '{TZ_1}',
                '- Техническое задание на создание дизайн-макетов отдельных блоков или страниц на сайте;'),

            'creating_pages': ('{CREATE_PAGES_1}', '- Создание страниц на сайте;'),
            'creating_pages2': ('{CREATE_PAGES_2}', '- Техническое задание на создание страниц на сайте;'),
            'creating_pages3': ('{CREATE_PAGES_3}', '')
        }

        contract_number = request.POST.get('contract_number')
        date_day = request.POST.get('date_day')
        site_name = request.POST.get('site_name')

        price_count_digit = request.POST.get('price_count_digit')
        price_count_word = num2words(price_count_digit, lang='ru')

        analitic_system = request.POST.getlist('analitic_system')
        analitic_system_user = request.POST.getlist('analitic_system_user')
        system_search = request.POST.getlist('search_system')

        date_month = request.POST.get('date_month')
        date_year = request.POST.get('date_year')
        organization_name = request.POST.get('organization_name')

        if organization_name.startswith('Индивидуальный предприниматель'):
            organization_name += ', именуемый'
        elif organization_name.startswith('ООО'):
            organization_name += ', именуемое'

        red_organization_name = request.POST.get('red_organization_name')

        guarantee = request.POST.get('guarantee')

        reason = request.POST.get('reason')
        person_name = request.POST.get('person_name')
        director_name = request.POST.get('director_name')
        email = request.POST.get('email')
        customer_id = request.POST.get('customer_id')
        inn = request.POST.get('inn')
        ogrn = request.POST.get('ogrn')
        registration_address = request.POST.get('registration_address')
        checking_account = request.POST.get('checking_account')
        correspondent_account = request.POST.get('correspondent_account')
        bank_name = request.POST.get('bank_name')
        bic = request.POST.get('bic')
        edo = request.POST.get('edo')
        choose_executor = request.POST.get('choose_executor')

        support_options = request.POST.getlist('support[]')
        platform_choice = request.POST.get('platform', None)
        selected_services = request.POST.getlist('services[]')

        template_filename = 'Договор Фикс метки.docx'
        template_path = os.path.join(os.path.dirname(__file__), '../dogovora', template_filename)
        doc = Document(template_path)
        signature_image_path = os.path.join(os.path.dirname(__file__), '../dogovora/podpis.jpg')

        if 'support_site' in support_options:
            support_text = '- Поддержка технического состояния сайта;'
            replace_tag_with_text(doc, '{SITE_SUPPORT}', support_text)
        else:
            replace_tag_with_text(doc, '{SITE_SUPPORT}')

        used_tags = set()

        for service_key in selected_services:
            if service_key in service_to_tag_mapping:
                tag, text = service_to_tag_mapping[service_key]
                replace_tag_with_text(doc, tag, text)
                used_tags.add(tag)

        all_tags = set(service_to_tag_mapping.values())
        unused_tags = {tag for tag, _ in all_tags if tag not in used_tags}
        for tag in unused_tags:
            replace_tag_with_text(doc, tag)

        handle_conditional_sections(doc, edo)
        handle_additional_work_sections(doc, platform_choice)

        only_yandex_selected = (
                analitic_system == ['yandex_web'] and
                analitic_system_user == ['yandex_metric'] and
                system_search == ['yandex_system']
        )

        if only_yandex_selected:
            # Изменение пункта 1.1
            replace_text_in_paragraphs(doc, 'в поисковых системах Яндекс и Google', 'в поисковой системе Яндекс')
            replace_text_in_paragraphs(doc, '(далее также – Системы)', '(далее также – Система)')

            # Изменение пункта 1.4.1
            remove_bullet_point(doc, '- Анализ ссылочного профиля сайта;')

            # Изменение пункта 1.4.4
            remove_bullet_point(doc, '- Работа с ссылочным профилем сайта (получение ссылок с других сайтов);')

            # Изменение пункта 1.4.6
            remove_bullet_point(doc,
                                '- Консультации и рекомендации по развитию сайта и продвижению в поисковых системах Яндекс и Google.')

            # Изменение пункта 2.1.1
            replace_text_in_paragraphs(doc, 'поисковых системах {YANDEX}{GOOGLE}', 'поисковой системе {YANDEX}{GOOGLE}')

        replace_analytics_tags(doc, analitic_system, analitic_system_user, system_search)
        if guarantee == 'with_guarantee':
            add_guarantee_section(doc)
        if choose_executor == 'ИП Михайлов Дмитрий Сергеевич':
            replace_tag_with_text(doc, '{PREDMET_DOGOVORA1}', 'по адаптации и оптимизации web-страниц')
            executor_name_replacement = ('Индивидуальный предприниматель Михайлов Дмитрий Сергеевич, именуемый в '
                                         'дальнейшем «Исполнитель», в лице генерального директора Михайлова Дмитрия '
                                         'Сергеевича, действующего '
                                         'на основании Свидетельства ОГРНИП 320784700136130')
            replacements_executor = {
                '{CHOOSE_EXECUTOR_NAME}': 'Индивидуальный предприниматель Михайлов Дмитрий Сергеевич',
                '{CHOOSE_EXECUTOR_INN}': '780256693210',
                '{CHOOSE_EXECUTOR_OGRNIP}': 'ОГРНИП: 320784700136130',
                '{CHOOSE_EXECUTOR_ADRESS}': '194295, Россия, г. Санкт-Петербург, пр-кт Северный, д. 24, корпус 1, '
                                            'кв. 33',
                '{CHOOSE_EXECUTOR_CHECKING_ACC}': '40802810201500152101',
                '{CHOOSE_EXECUTOR_KOR_ACC}': '30101810745374525104',
                '{CHOOSE_EXECUTOR_BANK}': 'ООО "Банк Точка"',
                '{CHOOSE_EXECUTOR_BIK}': '044525104',
                '{CHOOSE_EXECUTOR_EMAIL}': 'dima@mikhaylovseo.ru'
            }
        elif choose_executor == 'ООО «МД»':
            replace_tag_with_text(doc, '{PREDMET_DOGOVORA1}', 'по продвижению')
            executor_name_replacement = ('Общество с ограниченной ответственностью "Михайлов Диджитал", именуемое в '
                                         'дальнейшем «Исполнитель», в лице генерального директора Михайлова Дмитрия '
                                         'Сергеевича, действующего '
                                         'на основании Устава')
            replacements_executor = {
                '{CHOOSE_EXECUTOR_NAME}': 'Общество с ограниченной ответственностью "Михайлов Диджитал"',
                '{CHOOSE_EXECUTOR_INN}': '7810962062',
                '{CHOOSE_EXECUTOR_OGRNIP}': 'ОГРН: 1247800061464',
                '{CHOOSE_EXECUTOR_ADRESS}': '196142, Россия, г. Санкт-Петербург, ул. Пулковская, д. 2, корпус 1, '
                                            'литера А, оф 25, помещ. 66-Н',
                '{CHOOSE_EXECUTOR_CHECKING_ACC}': '40702810320000118082',
                '{CHOOSE_EXECUTOR_KOR_ACC}': '30101810745374525104',
                '{CHOOSE_EXECUTOR_BANK}': 'ООО "Банк Точка"',
                '{CHOOSE_EXECUTOR_BIK}': '044525104',
                '{CHOOSE_EXECUTOR_EMAIL}': 'dima@mikhaylovseo.ru'
            }
        replace_tag_with_text(doc, '{CHOOSE_EXECUTOR_NAME}', executor_name_replacement)

        for paragraph in doc.paragraphs:
            for key, value in replacements_executor.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements_executor.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)
                                for run in paragraph.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(9)

        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                for key, value in replacements_executor.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
                        for run in paragraph.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(9)

        footer = doc.sections[0].footer
        for paragraph in footer.paragraphs:
            paragraph.alignment = 1
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
        paragraph = footer.paragraphs[0]
        paragraph.add_run("________________" + director_name + "").font.size = Pt(12)
        paragraph.add_run("                                                                           ").font.size = Pt(
            12)

        if edo == "YES":
            paragraph.add_run("_______________Михайлов Д.С.").font.size = Pt(12)
        else:
            replace_underscores_with_signature(doc, "________________Михайлов Д.С.", signature_image_path)
            find_and_offset_director_text(doc)
            run = paragraph.add_run()
            run.add_picture(signature_image_path, width=Inches(0.8))  # Настройте ширину по необходимости
            paragraph.add_run("Михайлов Д.С.").font.size = Pt(12)

        add_newline_before_text(doc, "10.5. Договор составлен")
        replacements = {
            '{DOGOVOR_NUMBER}': contract_number,
            '{DAY}': date_day,
            '{MONTH}': date_month,
            '{YEAR}': date_year,
            '{CUSTOMER_ORGANIZATION}': organization_name,
            '{RED_CUSTOMER_ORGANIZATION}': red_organization_name,

            '{CUSTOMER_FIO}': person_name,
            '{DOGOVOR_OSNOVANIE}': reason,

            '{SITE_NAME}': site_name,
            '{PRICE_COUNT}': price_count_digit,
            '{PRICE_COUNT_IN_WORDS}': price_count_word,

            '{FIO_DIRECTOR}': director_name,
            '{CUSTOMER_EMAIL}': email,
            '{CUSTOMER_ID}': customer_id,
            '{CUSTOMER_NAME}': organization_name,
            '{INN}': inn,
            '{OGRN}': ogrn,
            '{REGISTRATION_ADDRESS}': registration_address,
            '{PAYMENT_ACCOUNT}': checking_account,
            '{CORRESPONDENT}': correspondent_account,
            '{BANK_NAME}': bank_name,
            '{BIK}': bic,
            '{CHOOSE_EXECUTOR}': choose_executor,
        }

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)
                        if key == '{DOGOVOR_NUMBER}' or key == '{CUSTOMER_NAME}':
                            run.bold = True

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)
                                for run in paragraph.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(9)
                                    if run.text.strip() == value:
                                        run.bold = True

        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9)

        text = organization_name
        make_text_bold_in_doc(doc, text)
        make_text_bold_in_doc(doc, 'Индивидуальный предприниматель Михайлов Дмитрий Сергеевич')
        make_text_bold_in_doc(doc, 'Общество с ограниченной ответственностью "Михайлов Диджитал"')

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="processed_contract.docx"'
        doc.save(response)
        return response
    else:
        return render(request, 'dogovor_FIX_form.html')
