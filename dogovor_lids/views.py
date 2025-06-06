import os
from django.http import HttpResponse
from django.shortcuts import render
from docx import Document
from docx.shared import Pt, Inches
from num2words import num2words

from dogovor_fix.views import replace_tag_with_text
from dogovor_yandex_direct.views import replace_tags_in_doc


def insert_edo_block(doc, edo: str, customer_id: str):
    if edo == 'YES':
        text = (
            "- ID Исполнителя в системе Тензор: 2BE894898d706174ab2aa3cdfc300550236\n"
            f"- ID Заказчика: {customer_id}\n"
            "9.4. Стороны согласовали, что они вправе осуществлять документооборот в электронном виде по телекоммуникационным каналам связи с использованием усиленной квалификационной электронной подписи посредством системы электронного документооборота СБИС.\n"
            "9.4.1. В целях настоящего договора под электронным документом понимается документ, созданный в электронной форме без предварительного документирования на бумажном носителе, подписанный электронной подписью в порядке, установленном законодательством Российской Федерации. Стороны признают электронные документы, заверенные электронной подписью, при соблюдении требований Федерального закона от 06.04.2011 № 63-ФЗ \"Об электронной подписи\" юридически эквивалентными документам на бумажных носителях, заверенным соответствующими подписями и оттиском печатей Сторон.\n"
            "9.5. Все изменения и дополнения к договору оформляются в виде дополнений и приложений к договору, являющихся его неотъемлемой частью.\n"
            "9.6. Договор составлен в двух подлинных экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из сторон, подписанных посредством ЭДО."
        )
    else:
        text = (
            "9.4. Все изменения и дополнения к договору оформляются в виде дополнений и приложений к договору, являющихся его неотъемлемой частью.\n"
            "9.5. Договор составлен в двух подлинных экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из сторон, подписанных лично."
        )

    replace_tag_with_text(doc, '{EDO_BLOCK}', text)


def insert_visitors_analytics(doc, yandex_metrica_enabled: bool):
    if yandex_metrica_enabled:
        text = '- Сбор и анализ данных о посетителях сайта в Яндекс. Метрика;'
    else:
        text = '- Сбор и анализ данных о посетителях сайта;'
    replace_tag_with_text(doc, '{VISITORS_ANALYTICS}', text)


def insert_support_section(doc, support_selected: bool):
    if support_selected:
        support_text = '- Поддержка технического состояния сайта;'
    else:
        support_text = ''
    replace_tag_with_text(doc, '{SUPPORT_SECTION}', support_text)


def insert_analytics_systems(doc, analytics_systems: list[str]):
    if 'yandex' in analytics_systems and 'search_console' in analytics_systems:
        analytics_text = 'Установка системы аналитики «Яндекс.Вебмастер» и «Search Console»;'
    elif 'yandex' in analytics_systems:
        analytics_text = 'Установка системы аналитики «Яндекс.Вебмастер»;'
    elif 'search_console' in analytics_systems:
        analytics_text = 'Установка системы аналитики «Search Console»;'
    else:
        analytics_text = ''

    replace_tag_with_text(doc, '{ANALYTICS_SYSTEMS}', analytics_text)


def insert_search_systems(doc, selected_search_systems: list[str]):
    # Формулировка с "далее – Система/Системы"
    if 'yandex' in selected_search_systems and 'google' in selected_search_systems:
        search_analytics = "в поисковых системах Яндекс и Google (далее также – Системы)"
        search_engines_full = "в поисковых системах Яндекс и Google"
    elif 'yandex' in selected_search_systems:
        search_analytics = "в поисковой системе Яндекс (далее также – Система)"
        search_engines_full = "в поисковой системе Яндекс"
    elif 'google' in selected_search_systems:
        search_analytics = "в поисковой системе Google (далее также – Система)"
        search_engines_full = "в поисковой системе Google"
    else:
        search_analytics = ""
        search_engines_full = ""

    replace_tag_with_text(doc, '{SEARCH_SYSTEMS}', search_analytics)
    replace_tag_with_text(doc, '{SEARCH_ENGINES_FULL}', search_engines_full)


def process_contract(request):
    if request.method == 'POST':
        contract_number = request.POST.get('contract_number')
        date_day = request.POST.get('date_day')
        date_month = request.POST.get('date_month')
        date_year = request.POST.get('date_year')
        organization_name = request.POST.get('organization_name')
        reason = request.POST.get('reason')
        person_name = request.POST.get('person_name')

        choose_executor = request.POST.get('choose_executor')
        site_name = request.POST.get('site_name')
        region_name = request.POST.get('region_name', 'Санкт-Петербург')
        custom_region_name = request.POST.get('custom_region_name', '').strip()

        if region_name == 'Другой' and custom_region_name:
            region_name = custom_region_name

        system_search = request.POST.getlist('search_system')

        optimization_map = {
            'headers': 'Оптимизация заголовков страниц;',
            'metatags': 'Оптимизация метатегов;',
            'texts': 'Написание текстов в помощь нейросетей и их оптимизация;',
            'structure': 'Оптимизация структуры сайта;',
            'tech_errors': 'Устранение технических ошибок;',
            'create_pages': 'Создание страниц на сайте;',
            'tz': 'Техническое задание (далее – ТЗ) на создание дизайн-макетов отдельных блоков или страниц на сайте;'
        }

        selected_opts = request.POST.getlist('opt')
        internal_optimization_lines = [
            f'- {optimization_map[opt]}' for opt in selected_opts if opt in optimization_map
        ]

        internal_optimization_text = '\n'.join(internal_optimization_lines)

        analytics_systems = request.POST.getlist('analytics_systems')

        support_selected = 'yes' in request.POST.getlist('support_site')
        yandex_metrica_enabled = 'yes' in request.POST.getlist('yandex_metrica')

        price_count = request.POST.get('price_count')
        price_count_in_words = num2words(price_count, lang='ru')

        customer_email = request.POST.get('email')
        customer_id = request.POST.get('customer_id')

        red_organization_name = request.POST.get('red_organization_name')

        director_name = request.POST.get('director_name')
        inn = request.POST.get('inn')
        ogrn = request.POST.get('ogrn')
        registration_address = request.POST.get('registration_address')
        checking_account = request.POST.get('checking_account')
        correspondent_account = request.POST.get('correspondent_account')
        bank_name = request.POST.get('bank_name')
        bic = request.POST.get('bic')
        edo = request.POST.get('edo')

        if choose_executor == 'ИП Михайлов Дмитрий Сергеевич':
            predmet_1 = 'адаптации и оптимизации web-страниц'
            predmet_2 = 'адаптации и оптимизации web-страниц'
            predmet_3 = 'адаптации и оптимизации web-страниц'

            replacements_executor = {
                '{CHOOSE_EXECUTOR_NAME}': 'Индивидуальный предприниматель Михайлов Дмитрий Сергеевич',
                '{CHOOSE_EXECUTOR_INN}': '780256693210',
                '{CHOOSE_EXECUTOR_OGRNIP}': 'ОГРНИП: 320784700136130',
                '{CHOOSE_EXECUTOR_ADRESS}': '194295, Россия, г. Санкт-Петербург, пр-кт Северный, д. 24, корпус 1, кв. 33',
                '{CHOOSE_EXECUTOR_CHECKING_ACC}': '40802810201500152101',
                '{CHOOSE_EXECUTOR_KOR_ACC}': '30101810745374525104',
                '{CHOOSE_EXECUTOR_BANK}': 'ООО "Банк Точка"',
                '{CHOOSE_EXECUTOR_BIK}': '044525104',
                '{CHOOSE_EXECUTOR_EMAIL}': 'info@mihaylov.digital'
            }

        else:
            predmet_1 = 'поисковой оптимизации'
            predmet_2 = 'поисковой оптимизации'
            predmet_3 = 'поисковой оптимизации'

            replacements_executor = {
                '{CHOOSE_EXECUTOR_NAME}': 'Общество с ограниченной ответственностью "Михайлов Диджитал"',
                '{CHOOSE_EXECUTOR_INN}': '7810962062',
                '{CHOOSE_EXECUTOR_OGRNIP}': 'ОГРН: 1247800061464',
                '{CHOOSE_EXECUTOR_ADRESS}': '196142, Россия, г. Санкт-Петербург, ул. Пулковская, д. 2, корпус 1, литера А, оф 25, помещ. 66-Н',
                '{CHOOSE_EXECUTOR_CHECKING_ACC}': '40702810320000118082',
                '{CHOOSE_EXECUTOR_KOR_ACC}': '30101810745374525104',
                '{CHOOSE_EXECUTOR_BANK}': 'ООО "Банк Точка"',
                '{CHOOSE_EXECUTOR_BIK}': '044525104',
                '{CHOOSE_EXECUTOR_EMAIL}': 'info@mihaylov.digital'
            }

            # Добавляем реквизиты исполнителя в общий список замен

        template_path = os.path.join(os.path.dirname(__file__), '../dogovora/Договор Лиды метки.docx')
        doc = Document(template_path)

        # Обработка вставки по ЭДО в пункт 10.3
        for paragraph in doc.paragraphs:
            if '{EDO}' in paragraph.text:
                if edo == 'NO':
                    lines = paragraph.text.split('\n')
                    cleaned_lines = []
                    for line in lines:
                        # Убираем строки с ЭДО
                        if not any(x in line for x in ['Либо посредством ЭДО', 'ID Исполнителя', 'ID Заказчика']):
                            cleaned_lines.append(line.replace('{EDO}', ''))
                    paragraph.text = '\n'.join(cleaned_lines)
                else:
                    paragraph.text = paragraph.text.replace('{EDO}', '')

                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9)

        replacements = {
            '{DOGOVOR_NUMBER}': contract_number,
            '{DAY}': date_day,
            '{MONTH}': date_month,
            '{YEAR}': date_year,
            '{CUSTOMER_ORGANIZATION}': organization_name,
            '{DOGOVOR_OSNOVANIE}': reason,
            '{CUSTOMER_FIO}': person_name,
            '{PREDMET_DOGOVORA1}': predmet_1,
            '{PREDMET_DOGOVORA2}': predmet_2,
            '{PREDMET_DOGOVORA3}': predmet_3,
            '{SITE_NAME}': site_name,
            '{PRICE_COUNT}': price_count,
            '{PRICE_COUNT_IN_WORDS}': price_count_in_words,
            '{REGION_NAME}': region_name,
            '{CUSTOMER_EMAIL}': customer_email,
            '{CUSTOMER_ID}': customer_id,

            '{RED_CUSTOMER_ORGANIZATION}': red_organization_name,

            '{FIO_DIRECTOR}': director_name,
            '{CUSTOMER_NAME}': organization_name,
            '{INN}': inn,
            '{OGRN}': ogrn,
            '{REGISTRATION_ADDRESS}': registration_address,
            '{PAYMENT_ACCOUNT}': checking_account,
            '{CORRESPONDENT}': correspondent_account,
            '{BANK_NAME}': bank_name,
            '{BIK}': bic,
            '{CHOOSE_EXECUTOR}': choose_executor,
        }
        replacements.update(replacements_executor)

        # Замена во всех параграфах
        replace_tags_in_doc(doc, replacements)
        insert_search_systems(doc, system_search)
        replace_tag_with_text(doc, '{INTERNAL_OPTIMIZATION}', internal_optimization_text)
        insert_analytics_systems(doc, analytics_systems)
        insert_support_section(doc, support_selected)
        insert_visitors_analytics(doc, yandex_metrica_enabled)
        insert_edo_block(doc, edo, customer_id)

        signature_image_path = os.path.join(os.path.dirname(__file__), '../dogovora/podpis.jpg')

        # Обработка колонтитула для подписи
        footer = doc.sections[0].footer
        footer.paragraphs.clear()  # очистим, чтобы не было наложений
        paragraph = footer.add_paragraph()
        paragraph.alignment = 1  # Центровка

        # Левая сторона – Заказчик
        run_left = paragraph.add_run("________________" + (director_name or ''))
        run_left.font.name = 'Calibri'
        run_left.font.size = Pt(9)

        # Пробелы между
        paragraph.add_run(" " * 30)  # достаточно, чтобы разделить подписи визуально

        # Правая сторона – Исполнитель
        if edo == "YES":
            run_exec = paragraph.add_run("_______________Михайлов Д.С.")
            run_exec.font.name = 'Calibri'
            run_exec.font.size = Pt(9)
        else:
            run_pic = paragraph.add_run()
            run_pic.add_picture(signature_image_path, width=Inches(0.8))
            run_text = paragraph.add_run("Михайлов Д.С.")
            run_text.font.name = 'Calibri'
            run_text.font.size = Pt(9)

        # Отдаём как docx-файл
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="contract_dogovor_lids.docx"'
        doc.save(response)
        return response

    # GET-запрос
    return render(request, '../templates/dogovor_lids_form.html')
