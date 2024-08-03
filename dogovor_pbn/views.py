import os
import re

from django.http import HttpResponse
from django.shortcuts import render

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_COLOR_INDEX
from docx.enum.table import WD_ALIGN_VERTICAL

from dogovor_fix.views import replace_tag_with_text

def make_text_bold_in_doc(doc, search_text):
    for paragraph in doc.paragraphs:
        if search_text in paragraph.text:
            runs = paragraph.runs
            for run in runs:
                if search_text in run.text:
                    split_text = run.text.split(search_text)
                    run.text = split_text[0]
                    bold_run = paragraph.add_run(search_text)
                    bold_run.bold = True
                    bold_run.font.name = 'Calibri'
                    bold_run.font.size = Pt(9)
                    if len(split_text) > 1:
                        after_bold_run = paragraph.add_run(split_text[1])
                        after_bold_run.font.name = 'Calibri'
                        after_bold_run.font.size = Pt(9)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if search_text in paragraph.text:
                        runs = paragraph.runs
                        for run in runs:
                            if search_text in run.text:
                                # Split the run text if it contains the search_text
                                split_text = run.text.split(search_text)
                                run.text = split_text[0]
                                bold_run = paragraph.add_run(search_text)
                                bold_run.bold = True
                                bold_run.font.name = 'Calibri'
                                bold_run.font.size = Pt(9)
                                if len(split_text) > 1:
                                    after_bold_run = paragraph.add_run(split_text[1])
                                    after_bold_run.font.name = 'Calibri'
                                    after_bold_run.font.size = Pt(9)


def remove_blank_line_between_points(doc, point1="3.1.", point2="3.2."):
    paragraphs = list(doc.paragraphs)
    for idx, paragraph in enumerate(paragraphs):
        if point1 in paragraph.text and idx + 2 < len(paragraphs):
            next_paragraph_text = paragraphs[idx + 1].text.strip()
            following_paragraph_text = paragraphs[idx + 2].text.strip()

            if next_paragraph_text == "" and point2 in following_paragraph_text:
                remove_paragraph(paragraphs[idx + 1])
                break


def remove_blank_line_between_points_2(doc, point1="1.1.", point2="1.2."):
    paragraphs = list(doc.paragraphs)
    for idx, paragraph in enumerate(paragraphs):
        if point1 in paragraph.text and idx + 2 < len(paragraphs):
            next_paragraph_text = paragraphs[idx + 1].text.strip()
            following_paragraph_text = paragraphs[idx + 2].text.strip()

            if next_paragraph_text == "" and point2 in following_paragraph_text:
                remove_paragraph(paragraphs[idx + 1])
                break


def adjust_paragraph_spacing(doc):
    paragraphs = list(doc.paragraphs)
    for idx, paragraph in enumerate(paragraphs):
        if "3.2." in paragraph.text and idx + 1 < len(paragraphs):
            next_paragraph_text = paragraphs[idx + 1].text.strip()

            if not next_paragraph_text or next_paragraph_text.startswith("4. Порядок приемки работ"):
                remove_paragraph(paragraphs[idx + 1])
                break


def remove_unnecessary_paragraphs(doc):
    for paragraph in doc.paragraphs:
        if "4. Анкоры и страницы" in paragraph.text:
            remove_paragraph(paragraph)
            break


def remove_paragraph(paragraph):
    p = paragraph._element
    if p.getparent() is not None:
        p.getparent().remove(p)
        p._p = p._element = None
def decrement_section_numbers(text):
    exclude_pattern = r'\d{2}\.\d{2}\.\d{4}|\d{2}-\w{1,3}-\w{1,3}'
    text = re.sub(exclude_pattern, lambda m: m.group().replace('.', '<dot>'), text)

    def replace(match):
        parts = match.group().split('.')
        if len(parts) >= 2 and parts[0].isdigit() and int(parts[0]) >= 5:
            parts[0] = str(int(parts[0]) - 1)
            return '.'.join(parts[:-1]) + '.'
        return match.group()

    pattern = r'\b\d+\.\d*\.?'
    updated_text = re.sub(pattern, replace, text)

    return updated_text.replace('<dot>', '.')


def replace_paragraph_text_with_styles(paragraph, new_text):
    """
    Заменяет текст в абзаце, сохраняя исходное форматирование.
    """
    if paragraph.runs:
        style = paragraph.runs[0].style
        font = paragraph.runs[0].font
        is_bold = font.bold
        font_name = font.name
        font_size = font.size
    else:
        is_bold = False
        font_name = 'Calibri'
        font_size = Pt(9)

    paragraph.clear()
    run = paragraph.add_run(new_text)
    run.bold = is_bold
    run.font.name = font_name
    run.font.size = font_size


def update_section_numbers(doc):
    start_update = False
    for paragraph in doc.paragraphs:
        if "5. Порядок приемки работ" in paragraph.text:
            start_update = True
        if start_update:
            updated_text = decrement_section_numbers(paragraph.text)
            replace_paragraph_text_with_styles(paragraph, updated_text)


def handle_conditional_sections(doc, predmet, site_creator, edo):
    edo_text_1 = "(в том числе его получения с использованием системы электронного документооборота)" if edo == "YES" else ""
    edo_text_2 = (
            "\nЛибо посредством ЭДО: "
            "\n- ID Исполнителя в системе Тензор: 2BE894898d706174ab2aa3cdfc300550236"
            "\n- ID Заказчика: {CUSTOMER_ID} "
            "\n10.4. Стороны согласовали, что они вправе осуществлять документооборот в электронном виде по телекоммуникационным каналам связи с использованием усиленной квалификационной электронной подписи посредством системы электронного документооборота. " + "\n" +
            "10.4.1. В целях настоящего договора под электронным документом понимается документ, созданный в электронной форме без предварительного документирования на бумажном носителе, подписанный электронной подписью в порядке, установленном законодательством Российской Федерации. Стороны признают электронные документы, заверенные электронной подписью, при соблюдении требований Федерального закона от 06.04.2011 № 63-ФЗ 'Об электронной подписи' юридически эквивалентными документам на бумажных носителях, заверенным соответствующими подписями и оттиском печатей Сторон. " + "\n" +
            "10.5. Все изменения и дополнения к договору оформляются в виде дополнений и приложений к договору, являющийся его неотъемлемой частью." +
            "\n10.6. Договор составлен в двух подлинных экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из сторон, подписанных лично либо посредством ЭДО. ") \
        if edo == "YES" else ""
    not_edo_text = "на почту Исполнителя" if edo == "NO" else ""

    write_by_hand = (
            "\n10.4. Все изменения и дополнения к договору оформляются в виде дополнений и приложений к договору, являющийся его неотъемлемой частью." +
            "\n10.5. Договор составлен в двух подлинных экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из сторон.") if edo == "NO" else ""

    replacements = {
        '{ARENDA_LINKS_1}': "",
        '{DROP_SEARCH_1}': "",
        '{ARENDA_LINKS_2}': "",
        '{DROP_SEARCH_2}': "",
        '{ARENDA_LINKS_3}': "",
        '{DROP_SEARCH_3}': "",
        '{ARENDA_LINKS_4}': "",
        '{DROP_SEARCH_4}': "",
        '{ARENDA_LINKS_5}': "",
        '{EDO_1}': edo_text_1,
        '{EDO_2}': edo_text_2,
        '{NOT_EDO}': not_edo_text,
        '{WRITE_BY_HAND}': write_by_hand
    }

    if predmet == "ARENDA_LINKS":
        replacements.update({
            '{ARENDA_LINKS_1}': "1.1. Исполнитель обязуется по заданию Заказчика оказать услуги по адаптации и оптимизации web-страниц сайтов своей площадки согласно техническому заданию, а Заказчик оплатить оказанные услуги.",
            '{ARENDA_LINKS_2}': "1.3.1. Написание 3 (трёх) околотематических статей без ссылок;\n1.3.2. Написание статьи по теме Заказчика или приближенной к ней и проставление ссылки с площадки Исполнителя.\n",
            '{ARENDA_LINKS_3}': "3.1. Стоимость аренды 1 (одной) ссылки на год на веб-сайтах площадки Исполнителя составляет 1 500 (тысяча пятьсот) рублей фиксированно. Количество ежемесячно закупаемых ссылок и суммарная оплата за них указана в Приложении №1 к Договору.",
            '{ARENDA_LINKS_4}': "4.1. Стороны признают целью Исполнителя – нахождение ссылок на страницы Интернет-сайта по необходимым анкорам на площадках Исполнителя. Анкоры и релевантные им страницы указаны в таблице в Приложении № 1 к Договору. Анкоры ссылок могут изменяться в рамках падежей, чисел, а также перестановкой слов и вставке до одного слова между словами или словосочетаниями.",
            '{ARENDA_LINKS_5}': ""})

    elif predmet == "DROP_SEARCH":
        hosting = "Исполнителя" if site_creator == "ISPOLNITEL" else "Заказчика"
        replacements.update({
            '{DROP_SEARCH_1}': f"1.1. Исполнитель обязуется по заданию Заказчика оказать услуги по поиску доменов и разработке веб-сайтов, а также их адаптации и модицификации согласно техническому заданию, а Заказчик оплатить оказанные услуги.",
            '{DROP_SEARCH_2}': f"1.3.1. Поиск доменов и их покупка на аккаунт Заказчика;\n1.3.2. Создание сайтов на WordPress со всеми плагинами на хостинге {hosting};\n1.3.3. Заказчик самостоятельно добавляет контент на сайт и ставит ссылки.\n",
            '{DROP_SEARCH_3}': "3.1. Оплата работ по адаптации и модификации веб-страниц сайтов площадки Исполнителя осуществляется Заказчиком ежемесячно, в порядке 100% предоплаты, в течение 5 (пяти) банковских дней, после выставления счета Исполнителем, если иное не оговорено в Дополнительном соглашении."
        })

    for paragraph in doc.paragraphs:
        for tag, replacement in replacements.items():
            if tag in paragraph.text:
                paragraph_text = paragraph.text.replace(tag, replacement)
                replace_paragraph_text_with_styles(paragraph, paragraph_text)


def set_cell_borders(cell):
    """Set borders for all sides of a cell."""
    # Define the borders
    sides = ('top', 'left', 'bottom', 'right')
    border_value = '6'  # A small border width value

    # Add border elements to the cell element
    tcPr = cell._tc.get_or_add_tcPr()
    for side in sides:
        tag = 'w:' + side
        element = parse_xml(r'<{0} {1} w:val="single" w:sz="{2}" '
                            r'w:space="0" w:color="auto"/>'.format(tag, nsdecls('w'), border_value))
        tcPr.append(element)


def replace_placeholder_with_image(doc, placeholder, image_path):
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.clear()
                    run.add_picture(image_path, width=Inches(0.8))
                    break


def replace_text_with_styles(paragraph, replacements_executor):
    for key, value in replacements_executor.items():
        if key in paragraph.text:
            runs = paragraph.runs
            full_text = ''.join(run.text for run in runs)
            split_text = full_text.split(key)
            for run in runs:
                run.text = ""
            paragraph.add_run(split_text[0])
            new_run = paragraph.add_run(value)
            new_run.bold = True  # Set the new text to bold
            new_run.font.name = 'Calibri'
            new_run.font.size = Pt(9)
            if len(split_text) > 1:
                paragraph.add_run(split_text[1])


def set_cell_formatting(cell, font_size=Pt(9), font_name='Calibri'):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in paragraph.runs:
            run.font.size = font_size
            run.font.name = font_name


def replace_underscores_with_signature(doc, placeholder_text, signature_path):
    def replace_in_paragraph(paragraph):
        runs = paragraph.runs
        full_text = ''.join(run.text for run in runs)
        if placeholder_text in full_text:
            split_text = full_text.split(placeholder_text)
            # Clear existing runs
            for run in runs:
                run.text = ""
            # Add new runs with the signature and styled text
            paragraph.add_run(split_text[0])
            paragraph.add_run().add_picture(signature_path, width=Inches(0.8))
            run = paragraph.add_run("Михайлов Д.С.")
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            if len(split_text) > 1:
                paragraph.add_run(split_text[1])

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            replace_in_paragraph(paragraph)
        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_in_paragraph(paragraph)

def find_and_offset_director_text(doc):
    for paragraph in doc.paragraphs:
        if "________________{FIO_DIRECTOR}" in paragraph.text:
            paragraph.text = "\n\n\n" + paragraph.text

    for section in doc.sections:
        for header in section.header.paragraphs:
            if "________________{FIO_DIRECTOR}" in header.text:
                header.text = "\n\n\n" + header.text
        for footer in section.footer.paragraphs:
            if "________________{FIO_DIRECTOR}" in footer.text:
                footer.text = "\n\n\n" + footer.text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "________________{FIO_DIRECTOR}" in paragraph.text:
                        paragraph.text = "\n\n\n" + paragraph.text


def process_contract(request):
    if request.method == 'POST':
        contract_number = request.POST.get('contract_number')
        date_day = request.POST.get('date_day')
        date_month = request.POST.get('date_month')
        date_year = request.POST.get('date_year')
        organization_name = request.POST.get('organization_name')

        if organization_name.startswith('Индивидуальный предприниматель'):
            organization_name += ', именуемый'
        elif organization_name.startswith('ООО'):
            organization_name += ', именуемое'

        red_organization_name = request.POST.get('red_organization_name')
        reason = request.POST.get('reason')
        person_name = request.POST.get('person_name')
        director_name = request.POST.get('director_name')
        month_count = request.POST.get('month_count')
        email = request.POST.get('email')
        customer_id = request.POST.get('customer_id')

        choose_executor = request.POST.get('choose_executor')

        inn = request.POST.get('inn')
        ogrn = request.POST.get('ogrn')
        registration_address = request.POST.get('registration_address')
        checking_account = request.POST.get('checking_account')
        correspondent_account = request.POST.get('correspondent_account')
        bank_name = request.POST.get('bank_name')
        bic = request.POST.get('bic')
        predmet = request.POST.get('predmet')
        edo = request.POST.get('edo')
        site_creation = request.POST.get('site_creation')

        template_filename = 'Договор PBN_динамика.docx'
        template_path = os.path.join(os.path.dirname(__file__), '../dogovora', template_filename)
        doc = Document(template_path)
        handle_conditional_sections(doc, predmet, site_creation, edo)

        if choose_executor == 'ИП Михайлов Дмитрий Сергеевич':
            executor_name_replacement = ('Индивидуальный предприниматель Михайлов Дмитрий Сергеевич, именуемый в '
                                         'дальнейшем «Исполнитель», в лице генерального директора Михайлова Дмитрия '
                                         'Сергеевича, действующего'
                                         'на основании Свидетельства ОГРНИП 320784700136130')
            replacements_executor = {
                '{CHOOSE_EXECUTOR_NAME}': 'Индивидуальный предприниматель Михайлов Дмитрий Сергеевич',
                '{CHOOSE_EXECUTOR_INN}': '780256693210',
                '{CHOOSE_EXECUTOR_OGRNIP}': 'ОГРНИП: 320784700136130',
                '{CHOOSE_EXECUTOR_ADRESS}': '194295, Россия, г. Санкт-Петербург, пр-кт Северный, д. 24, корпус 1, '
                                            'кв. 33',
                '{CHOOSE_EXECUTOR_CHECKING_ACC}': '40802810201500152101',
                '{CHOOSE_EXECUTOR_KOR_ACC}': '30101810745374525104',
                '{CHOOSE_EXECUTOR_BANK}': 'ООО "Банк Точка"',
                '{CHOOSE_EXECUTOR_BIK}': '044525104',
                '{CHOOSE_EXECUTOR_EMAIL}': 'dima@mikhaylovseo.ru'
            }
        elif choose_executor == 'ООО «МД»':
            executor_name_replacement = ('Общество с ограниченной ответственностью "Михайлов Диджитал", именуемый в '
                                         'дальнейшем «Исполнитель», в лице генерального директора Михайлова Дмитрия '
                                         'Сергеевича, действующего'
                                         'на основании Устава')
            replacements_executor = {
                '{CHOOSE_EXECUTOR_NAME}': 'Общество с ограниченной ответственностью "Михайлов Диджитал"',
                '{CHOOSE_EXECUTOR_INN}': '7810962062',
                '{CHOOSE_EXECUTOR_OGRNIP}': 'ОГРН: 1247800061464',
                '{CHOOSE_EXECUTOR_ADRESS}': '196142, Россия, г. Санкт-Петербург, ул. Пулковская, д. 2, корпус 1, '
                                            'литера А, оф 25, помещ. 66-Н',
                '{CHOOSE_EXECUTOR_CHECKING_ACC}': '40702810320000118082',
                '{CHOOSE_EXECUTOR_KOR_ACC}': '30101810745374525104',
                '{CHOOSE_EXECUTOR_BANK}': 'ООО "Банк Точка"',
                '{CHOOSE_EXECUTOR_BIK}': '044525104',
                '{CHOOSE_EXECUTOR_EMAIL}': 'dima@mikhaylovseo.ru'
            }

        replace_tag_with_text(doc, '{CHOOSE_EXECUTOR_NAME}', executor_name_replacement)

        for paragraph in doc.paragraphs:
            for key, value in replacements_executor.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements_executor.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)
                                for run in paragraph.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(9)

        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                for key, value in replacements_executor.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
                        for run in paragraph.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(9)

        signature_image_path = os.path.join(os.path.dirname(__file__), '../dogovora/podpis.jpg')
        if predmet == "DROP_SEARCH":
            remove_unnecessary_paragraphs(doc)
            update_section_numbers(doc)
            adjust_paragraph_spacing(doc)

        footer = doc.sections[0].footer
        for paragraph in footer.paragraphs:
            paragraph.alignment = 1
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9)

        paragraph = footer.paragraphs[0]
        paragraph.add_run("________________" + director_name + "").font.size = Pt(12)
        paragraph.add_run("                                                                           ").font.size = Pt(
            12)

        if edo == "YES":
            paragraph.add_run("_______________Михайлов Д.С.").font.size = Pt(12)
            remove_blank_line_between_points(doc)
        else:
            replace_underscores_with_signature(doc, "_______________Михайлов Д.С.", signature_image_path)
            find_and_offset_director_text(doc)
            run = paragraph.add_run()
            run.add_picture(signature_image_path, width=Inches(0.8))  # Настройте ширину по необходимости
            paragraph.add_run("Михайлов Д.С.").font.size = Pt(12)
            remove_blank_line_between_points(doc)
            remove_blank_line_between_points_2(doc)

        replacements = {
            '{DOGOVOR_NUMBER}': contract_number,
            '{DAY}': date_day,
            '{MONTH}': date_month,
            '{YEAR}': date_year,
            '{CUSTOMER_ORGANIZATION}': organization_name,
            '{RED_CUSTOMER_ORGANIZATION}': red_organization_name,
            '{CUSTOMER_FIO}': person_name,
            '{DOGOVOR_OSNOVANIE}': reason,
            '{FIO_DIRECTOR}': director_name,
            '{CUSTOMER_EMAIL}': email,
            '{CUSTOMER_ID}': customer_id,
            '{CUSTOMER_NAME}': 'ИП ' + person_name,
            '{INN}': inn,
            '{OGRN}': ogrn,
            '{REGISTRATION_ADDRESS}': registration_address,
            '{PAYMENT_ACCOUNT}': checking_account,
            '{CORRESPONDENT}': correspondent_account,
            '{BANK_NAME}': bank_name,
            '{BIK}': bic,
            '{MONTH_COUNT}': month_count,
            '{CHOOSE_EXECUTOR}': choose_executor,
        }

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)
                        if key == '{DOGOVOR_NUMBER}' or key == '{CUSTOMER_NAME}':
                            run.bold = True

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)
                                for run in paragraph.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(9)
                                    if run.text.strip() == value:
                                        run.bold = True

        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9)

        text = organization_name
        make_text_bold_in_doc(doc, text)
        make_text_bold_in_doc(doc,'Индивидуальный предприниматель Михайлов Дмитрий Сергеевич')
        make_text_bold_in_doc(doc,'Общество с ограниченной ответственностью "Михайлов Диджитал"')

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="processed_contract.docx"'
        doc.save(response)
        return response
    else:
        return render(request, 'dogovor_PBN_form.html')
