import os
from django.http import HttpResponse
from django.shortcuts import render
from docx import Document
from docx.shared import Pt, Inches
from num2words import num2words


def make_text_bold_in_doc(doc, search_text):
    main_text = search_text.split(", именуемое")[0]  # Основная часть текста до ", именуемое"

    for paragraph in doc.paragraphs:
        if main_text in paragraph.text:
            runs = paragraph.runs
            for run in runs:
                if main_text in run.text:
                    split_text = run.text.split(main_text, 1)
                    run.text = split_text[0]
                    bold_run = paragraph.add_run(main_text)
                    bold_run.bold = True
                    bold_run.font.name = 'Calibri'
                    bold_run.font.size = Pt(9)
                    if len(split_text) > 1:
                        after_bold_run = paragraph.add_run(split_text[1])
                        after_bold_run.font.name = 'Calibri'
                        after_bold_run.font.size = Pt(9)

    # Повторение для таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if main_text in paragraph.text:
                        runs = paragraph.runs
                        for run in runs:
                            if main_text in run.text:
                                split_text = run.text.split(main_text, 1)
                                run.text = split_text[0]
                                bold_run = paragraph.add_run(main_text)
                                bold_run.bold = True
                                bold_run.font.name = 'Calibri'
                                bold_run.font.size = Pt(9)
                                if len(split_text) > 1:
                                    after_bold_run = paragraph.add_run(split_text[1])
                                    after_bold_run.font.name = 'Calibri'
                                    after_bold_run.font.size = Pt(9)


def replace_paragraph_text_with_styles(paragraph, new_text):
    """
    Заменяет текст в абзаце, сохраняя исходное форматирование.
    """
    if paragraph.runs:
        style = paragraph.runs[0].style
        font = paragraph.runs[0].font
        is_bold = font.bold
        font_name = font.name
        font_size = font.size
    else:
        is_bold = False
        font_name = 'Calibri'
        font_size = Pt(9)

    paragraph.clear()
    run = paragraph.add_run(new_text)
    run.bold = is_bold
    run.font.name = font_name
    run.font.size = font_size


def handle_conditional_sections(doc, edo):
    edo_text_1 = "(в том числе его получения с использованием системы электронного документооборота)" if edo == "YES" else ""
    edo_text_2 = (
            "\nЛибо посредством ЭДО: "
            "\n- ID Исполнителя в системе Тензор: 2BE894898d706174ab2aa3cdfc300550236"
            "\n- ID Заказчика: {CUSTOMER_ID} "
            "\n10.4. Стороны согласовали, что они вправе осуществлять документооборот в электронном виде по телекоммуникационным каналам связи с использованием усиленной квалификационной электронной подписи посредством системы электронного документооборота. " + "\n" +
            "10.4.1. В целях настоящего договора под электронным документом понимается документ, созданный в электронной форме без предварительного документирования на бумажном носителе, подписанный электронной подписью в порядке, установленном законодательством Российской Федерации. Стороны признают электронные документы, заверенные электронной подписью, при соблюдении требований Федерального закона от 06.04.2011 № 63-ФЗ 'Об электронной подписи' юридически эквивалентными документам на бумажных носителях, заверенным соответствующими подписями и оттиском печатей Сторон." + "\n" +
            "10.5. Все изменения и дополнения к договору оформляются в виде дополнений и приложений к договору, являющийся его неотъемлемой частью." + "\n" +
            "10.6. Договор составлен в двух подлинных экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из сторон, подписанных лично либо посредством ЭДО. ") \
        if edo == "YES" else ""
    not_edo_text = "на почту Исполнителя" if edo == "NO" else ""

    write_by_hand = (
            "\n10.4. Все изменения и дополнения к договору оформляются в виде дополнений и приложений к договору, являющийся его неотъемлемой частью. " + '\n' +
            "10.5. Договор составлен в двух подлинных экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из сторон.") if edo == "NO" else ""

    replacements = {
        '{EDO_1}': edo_text_1,
        '{EDO_2}': edo_text_2,
        '{NOT_EDO}': not_edo_text,
        '{WRITE_BY_HAND}': write_by_hand
    }

    for paragraph in doc.paragraphs:
        for tag, replacement in replacements.items():
            if tag in paragraph.text:
                paragraph_text = paragraph.text.replace(tag, replacement)
                replace_paragraph_text_with_styles(paragraph, paragraph_text)


def replace_tag_with_text(doc, tag, text=None):
    for paragraph in doc.paragraphs:
        if tag in paragraph.text:
            new_text = paragraph.text.replace(tag, text if text is not None else '').strip()
            if new_text:
                replace_paragraph_text_with_styles(paragraph, new_text)
            else:
                # Если параграф остался пустым, удаляем его
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None


def replace_underscores_with_signature(doc, placeholder_text, signature_path):
    def replace_in_paragraph(paragraph):
        runs = paragraph.runs
        full_text = ''.join(run.text for run in runs)
        if placeholder_text in full_text:
            split_text = full_text.split(placeholder_text)
            # Clear existing runs
            for run in runs:
                run.text = ""
            # Add new runs with the signature and styled text
            paragraph.add_run(split_text[0])
            paragraph.add_run().add_picture(signature_path, width=Inches(0.8))
            run = paragraph.add_run("Михайлов Д.С.")
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            if len(split_text) > 1:
                paragraph.add_run(split_text[1])

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            replace_in_paragraph(paragraph)
        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_in_paragraph(paragraph)


def find_and_offset_director_text(doc):
    # Обход всех параграфов
    for paragraph in doc.paragraphs:
        if "________________{FIO_DIRECTOR}" in paragraph.text:
            paragraph.text = "\n\n\n" + paragraph.text

    # Обход всех секций (header и footer)
    for section in doc.sections:
        for header in section.header.paragraphs:
            if "________________{FIO_DIRECTOR}" in header.text:
                header.text = "\n\n\n" + header.text
        for footer in section.footer.paragraphs:
            if "________________{FIO_DIRECTOR}" in footer.text:
                footer.text = "\n\n\n" + footer.text

    # Обход всех таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "________________{FIO_DIRECTOR}" in paragraph.text:
                        paragraph.text = "\n\n\n" + paragraph.text


def process_contract(request):
    if request.method == 'POST':

        contract_number = request.POST.get('contract_number')
        date_day = request.POST.get('date_day')
        site_name = request.POST.get('site_name')

        price_count = request.POST.get('price_count')
        price_count_in_words = num2words(price_count, lang='ru')

        link = request.POST.get('link')

        date_month = request.POST.get('date_month')
        date_year = request.POST.get('date_year')

        organization_name = request.POST.get('organization_name')

        if organization_name.startswith('Индивидуальный предприниматель'):
            organization_name += ', именуемый'
        elif organization_name.startswith('ООО'):
            organization_name += ', именуемое'

        red_organization_name = request.POST.get('red_organization_name')
        customer_id = request.POST.get('customer_id')

        reason = request.POST.get('reason')
        person_name = request.POST.get('person_name')
        director_name = request.POST.get('director_name')
        email = request.POST.get('email')
        inn = request.POST.get('inn')
        ogrn = request.POST.get('ogrn')
        registration_address = request.POST.get('registration_address')
        checking_account = request.POST.get('checking_account')
        correspondent_account = request.POST.get('correspondent_account')
        bank_name = request.POST.get('bank_name')
        bic = request.POST.get('bic')
        edo = request.POST.get('edo')
        choose_executor = request.POST.get('choose_executor')

        template_filename = 'Договор ПФ метки.docx'
        template_path = os.path.join(os.path.dirname(__file__), '../dogovora', template_filename)
        doc = Document(template_path)
        signature_image_path = os.path.join(os.path.dirname(__file__), '../dogovora/podpis.jpg')

        handle_conditional_sections(doc, edo)

        if choose_executor == 'ИП Михайлов Дмитрий Сергеевич':
            replace_tag_with_text(doc, '{PREDMET_DOGOVORA1}', 'услуги по адаптации и оптимизации web-страниц')
            executor_name_replacement = ('Индивидуальный предприниматель Михайлов Дмитрий Сергеевич, именуемый в '
                                         'дальнейшем «Исполнитель», в лице Михайлова Дмитрия '
                                         'Сергеевича, действующего'
                                         ' на основании Свидетельства ОГРНИП 320784700136130')
            replace_tag_with_text(doc, '{ADAPTATION}', 'адаптации и оптимизации web-страниц')
            replacements_executor = {
                '{CHOOSE_EXECUTOR_NAME}': 'Индивидуальный предприниматель Михайлов Дмитрий Сергеевич',
                '{CHOOSE_EXECUTOR_INN}': '780256693210',
                '{CHOOSE_EXECUTOR_OGRNIP}': 'ОГРНИП: 320784700136130',
                '{CHOOSE_EXECUTOR_ADRESS}': '194295, Россия, г. Санкт-Петербург, пр-кт Северный, д. 24, корпус 1, '
                                            'кв. 33',
                '{CHOOSE_EXECUTOR_CHECKING_ACC}': '40802810201500152101',
                '{CHOOSE_EXECUTOR_KOR_ACC}': '30101810745374525104',
                '{CHOOSE_EXECUTOR_BANK}': 'ООО "Банк Точка"',
                '{CHOOSE_EXECUTOR_BIK}': '044525104',
                '{CHOOSE_EXECUTOR_EMAIL}': 'info@mihaylov.digital'
            }
        elif choose_executor == 'ООО «МД»':
            replace_tag_with_text(doc, '{PREDMET_DOGOVORA1}', 'рекламные услуги по поисковой оптимизации')
            executor_name_replacement = ('Общество с ограниченной ответственностью "Михайлов Диджитал", именуемый в '
                                         'дальнейшем «Исполнитель», в лице генерального директора Михайлова Дмитрия '
                                         'Сергеевича, действующего '
                                         'на основании Устава')
            replace_tag_with_text(doc, '{ADAPTATION}', 'поисковой оптимизации')
            replacements_executor = {
                '{CHOOSE_EXECUTOR_NAME}': 'Общество с ограниченной ответственностью "Михайлов Диджитал"',
                '{CHOOSE_EXECUTOR_INN}': '7810962062',
                '{CHOOSE_EXECUTOR_OGRNIP}': 'ОГРН: 1247800061464',
                '{CHOOSE_EXECUTOR_ADRESS}': '196142, Россия, г. Санкт-Петербург, ул. Пулковская, д. 2, корпус 1, '
                                            'литера А, оф 25, помещ. 66-Н',
                '{CHOOSE_EXECUTOR_CHECKING_ACC}': '40702810320000118082',
                '{CHOOSE_EXECUTOR_KOR_ACC}': '30101810745374525104',
                '{CHOOSE_EXECUTOR_BANK}': 'ООО "Банк Точка"',
                '{CHOOSE_EXECUTOR_BIK}': '044525104',
                '{CHOOSE_EXECUTOR_EMAIL}': 'info@mihaylov.digital'
            }
        replace_tag_with_text(doc, '{CHOOSE_EXECUTOR_NAME}', executor_name_replacement)

        for paragraph in doc.paragraphs:
            for key, value in replacements_executor.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements_executor.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)
                                for run in paragraph.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(9)

        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                for key, value in replacements_executor.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
                        for run in paragraph.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(9)

        footer = doc.sections[0].footer
        for paragraph in footer.paragraphs:
            paragraph.alignment = 1
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
        paragraph = footer.paragraphs[0]
        paragraph.add_run("________________" + director_name + "").font.size = Pt(12)
        paragraph.add_run("                                                                           ").font.size = Pt(
            12)

        if edo == "YES":
            paragraph.add_run("_______________Михайлов Д.С.").font.size = Pt(12)
        else:
            replace_underscores_with_signature(doc, "________________Михайлов Д.С.", signature_image_path)
            find_and_offset_director_text(doc)
            run = paragraph.add_run()
            run.add_picture(signature_image_path, width=Inches(0.8))
            paragraph.add_run("Михайлов Д.С.").font.size = Pt(12)

        replacements = {
            '{DOGOVOR_NUMBER}': contract_number,
            '{DAY}': date_day,
            '{MONTH}': date_month,
            '{YEAR}': date_year,
            '{CUSTOMER_ORGANIZATION}': organization_name,
            '{RED_CUSTOMER_ORGANIZATION}': red_organization_name,

            '{CUSTOMER_FIO}': person_name,
            '{DOGOVOR_OSNOVANIE}': reason,

            '{SITE_NAME}': site_name,
            '{PRICE_COUNT}': price_count,
            '{PRICE_COUNT_IN_WORDS}': price_count_in_words,
            '{LINK}': link,
            '{CUSTOMER_ID}': customer_id,

            '{FIO_DIRECTOR}': director_name,
            '{CUSTOMER_EMAIL}': email,
            '{CUSTOMER_NAME}': organization_name,
            '{INN}': inn,
            '{OGRN}': ogrn,
            '{REGISTRATION_ADDRESS}': registration_address,
            '{PAYMENT_ACCOUNT}': checking_account,
            '{CORRESPONDENT}': correspondent_account,
            '{BANK_NAME}': bank_name,
            '{BIK}': bic,
            '{CHOOSE_EXECUTOR}': choose_executor,
        }

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)
                        if key == '{DOGOVOR_NUMBER}' or key == '{CUSTOMER_NAME}':
                            run.bold = True

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)
                                for run in paragraph.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(9)
                                    if run.text.strip() == value:
                                        run.bold = True

        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9)

        text = organization_name
        make_text_bold_in_doc(doc, text)
        make_text_bold_in_doc(doc, 'Индивидуальный предприниматель Михайлов Дмитрий Сергеевич')
        make_text_bold_in_doc(doc, 'Общество с ограниченной ответственностью "Михайлов Диджитал"')

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="processed_contract.docx"'
        doc.save(response)
        return response
    else:
        return render(request, 'dogovor_PF_form.html')
