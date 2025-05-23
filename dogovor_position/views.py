import os
from django.http import HttpResponse
from django.shortcuts import render
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches
from docx.text.paragraph import Paragraph
from num2words import num2words


def replace_text_in_paragraphs(doc, search_text, replace_text):
    for paragraph in doc.paragraphs:
        if search_text in paragraph.text:
            new_text = paragraph.text.replace(search_text, replace_text)
            replace_paragraph_text_with_styles(paragraph, new_text)


def remove_bullet_point(doc, bullet_text):
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == bullet_text:
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None
            break  # Предполагаем, что такой пункт единственный


def insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_paragraph.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
    if style:
        new_paragraph.style = style
    return new_paragraph


def add_guarantee_section(doc):
    # Флаг, который указывает, что мы находимся внутри раздела 4
    in_section_4 = False

    # Текст новых пунктов
    new_section_header = '4. Ключевые фразы и гарантия'
    new_point_4_4 = '4.4. Гарантией выполненных работ является вывод в топ-10 поисковой сети Яндекс не менее 30% запросов, указанных в Приложении №2, в течение 6 (шести) месяцев работ.'
    new_point_4_5 = '4.5. При невыполнении гарантийных обязательств, Исполнитель предоставляет 7 (седьмой) месяц работ бесплатно.'
    new_point_4_6 = '''4.6. Условия гарантии являются недействительными в случаях:
- действий или бездействий поисковых систем, хостинга, системы сайта – для этого важно настраивать журнал логов;
- форс-мажорных обстоятельств, указанных в п. 8 настоящего Договора;
- работ Заказчика и третьих лиц на сайте без согласования с Исполнителем;
- смены стратегии по инициативе Заказчика;
- несоблюдения/несвоевременного соблюдения выполнения Заказчиком данных Исполнителем рекомендаций по п.1.4. настоящего Договора.'''

    paragraphs = doc.paragraphs
    i = 0
    while i < len(paragraphs):
        paragraph = paragraphs[i]
        text = paragraph.text.strip()

        if text.startswith('4. '):
            # Находим заголовок раздела 4 и заменяем его
            paragraph.clear()
            run = paragraph.add_run(new_section_header)
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            run.bold = True
            in_section_4 = True
            i += 1
            continue

        if in_section_4:
            if text.startswith('5. '):
                # Дошли до раздела 5, выходим из раздела 4
                in_section_4 = False
                i += 1
                continue

            if text.startswith('4.4. '):
                # Заменяем пункт 4.4 и добавляем новые пункты
                paragraph.clear()
                run = paragraph.add_run(new_point_4_4)
                run.font.name = 'Calibri'
                run.font.size = Pt(9)

                # Вставляем новый параграф для 4.5
                p4_5 = insert_paragraph_after(paragraph, new_point_4_5)
                # Вставляем параграфы для 4.6 и его подпунктов
                lines = new_point_4_6.split('\n')
                previous_p = p4_5
                for line in lines:
                    p = insert_paragraph_after(previous_p, line)
                    previous_p = p  # Обновляем ссылку на предыдущий параграф

                # После добавления новых пунктов, пропускаем оригинальный 4.5 и 4.6, если они есть
                # Проверяем следующие параграфы
                while i + 1 < len(paragraphs) and paragraphs[i + 1].text.strip().startswith('4.'):
                    i += 1
                i += 1
                continue

            else:
                i += 1
                continue
        else:
            i += 1


def make_text_bold_in_doc(doc, search_text):
    main_text = search_text.split(", именуемое")[0]  # Основная часть текста до ", именуемое"

    for paragraph in doc.paragraphs:
        if main_text in paragraph.text:
            runs = paragraph.runs
            for run in runs:
                if main_text in run.text:
                    split_text = run.text.split(main_text, 1)
                    run.text = split_text[0]
                    bold_run = paragraph.add_run(main_text)
                    bold_run.bold = True
                    bold_run.font.name = 'Calibri'
                    bold_run.font.size = Pt(9)
                    if len(split_text) > 1:
                        after_bold_run = paragraph.add_run(split_text[1])
                        after_bold_run.font.name = 'Calibri'
                        after_bold_run.font.size = Pt(9)

    # Повторение для таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if main_text in paragraph.text:
                        runs = paragraph.runs
                        for run in runs:
                            if main_text in run.text:
                                split_text = run.text.split(main_text, 1)
                                run.text = split_text[0]
                                bold_run = paragraph.add_run(main_text)
                                bold_run.bold = True
                                bold_run.font.name = 'Calibri'
                                bold_run.font.size = Pt(9)
                                if len(split_text) > 1:
                                    after_bold_run = paragraph.add_run(split_text[1])
                                    after_bold_run.font.name = 'Calibri'
                                    after_bold_run.font.size = Pt(9)


def replace_paragraph_text_with_styles(paragraph, new_text):
    """
    Заменяет текст в абзаце, сохраняя исходное форматирование.
    """
    if paragraph.runs:
        style = paragraph.runs[0].style
        font = paragraph.runs[0].font
        is_bold = font.bold
        font_name = font.name
        font_size = font.size
    else:
        is_bold = False
        font_name = 'Calibri'
        font_size = Pt(9)

    paragraph.clear()
    run = paragraph.add_run(new_text)
    run.bold = is_bold
    run.font.name = font_name
    run.font.size = font_size


def handle_additional_work_sections(doc, selected_services, platform_choice):
    # Найти начальный и конечный индексы параграфов для удаления
    start_index, end_index = None, None
    for i, paragraph in enumerate(doc.paragraphs):
        if '{IF_TZ_DELETE}' in paragraph.text:
            start_index = i
        elif '{END_OF_SECTION}' in paragraph.text:
            end_index = i
            break

    if start_index is not None and end_index is not None:
        if 'design_layouts' in selected_services:
            # Удаляем параграфы между start_index и end_index включительно
            for i in range(start_index, end_index + 1)[::-1]:  # В обратном порядке
                p = doc.paragraphs[i]._element
                p.getparent().remove(p)
                p._p = p._element = None
        else:
            for i in [start_index, end_index]:
                paragraph = doc.paragraphs[i]
                text = paragraph.text.replace('{IF_TZ_DELETE}', '').replace('{END_OF_SECTION}', '')
                replace_paragraph_text_with_styles(paragraph, text)

    for paragraph in doc.paragraphs:
        if '{WORD_PRESS}' in paragraph.text or '{NOT_WORD_PRESS}' in paragraph.text:
            if platform_choice == 'wordpress':
                text = paragraph.text.replace('{WORD_PRESS}',
                                              ' по результатам коммерческого аудита. Работы программиста после проведения других аудитов включены в счёт.').replace(
                    '{NOT_WORD_PRESS}', '')
            elif platform_choice == 'not_wordpress':
                text = paragraph.text.replace('{WORD_PRESS}', '').replace('{NOT_WORD_PRESS}',
                                                                          '.')
            else:
                text = paragraph.text.replace('{WORD_PRESS}', '').replace('{NOT_WORD_PRESS}', '')
            replace_paragraph_text_with_styles(paragraph, text)


def calculate_and_replace_tags_10(doc, request):
    req_count_top10 = request.POST.get('req_count_top10')
    if req_count_top10:
        req_count_top10 = int(req_count_top10)
        region_name = request.POST.get('region_name')
        search_engine = request.POST.get('search_engine')

        google_req_cost = 0
        if search_engine in ['GOOGLE', 'YANDEX_GOOGLE']:
            google_req_cost = int(request.POST.get('google_req_cost'))

        if region_name == 'Москва':
            req_pay = int(request.POST.get('req_count_top10_msk'))
        elif region_name == 'Санкт-Петербург':
            req_pay = int(request.POST.get('req_count_top10_spb'))
        else:
            req_pay = int(request.POST.get('req_count_top10_other'))

        req_pay += google_req_cost

        replace_tag_with_text(doc, '{REQ_PAY_10}', str(req_pay))
        replace_tag_with_text(doc, '{REQ_PAY_WORDS_10}', num2words(req_pay, lang='ru'))

        total_premium = req_pay * req_count_top10
        for i in range(10, 101, 10):
            percent = i / 100
            tag = '{PREM_TOP10_' + str(i) + '}'
            prem_value = int(total_premium * percent)
            replace_tag_with_text(doc, tag, str(prem_value))
    else:
        replace_tag_with_text(doc, '{REQ_PAY_10}', "__________")
        replace_tag_with_text(doc, '{REQ_COUNT_TOP10}', "__________")
        replace_tag_with_text(doc, '{REQ_PAY_WORDS_10}', "__________")
        for i in range(10, 101, 10):
            tag = '{PREM_TOP10_' + str(i) + '}'
            replace_tag_with_text(doc, tag, "__________")


def calculate_and_replace_tags_5(doc, request):
    req_count_top5 = request.POST.get('req_count_top5')
    if req_count_top5:
        req_count_top5 = int(req_count_top5)
        region_name = request.POST.get('region_name')
        search_engine = request.POST.get('search_engine')

        google_req_cost = 0
        if search_engine in ['GOOGLE', 'YANDEX_GOOGLE']:
            google_req_cost = int(request.POST.get('google_req_cost'))

        if region_name == 'Москва':
            req_pay = int(request.POST.get('req_count_top5_msk'))
        elif region_name == 'Санкт-Петербург':
            req_pay = int(request.POST.get('req_count_top5_spb'))
        else:
            req_pay = int(request.POST.get('req_count_top5_other'))

        req_pay += google_req_cost

        replace_tag_with_text(doc, '{REQ_PAY_5}', str(req_pay))
        replace_tag_with_text(doc, '{REQ_PAY_WORDS_5}', num2words(req_pay, lang='ru'))

        total_premium = req_pay * req_count_top5
        for i in range(10, 101, 10):
            percent = i / 100
            tag = '{PREM_TOP5_' + str(i) + '}'
            prem_value = int(total_premium * percent)
            replace_tag_with_text(doc, tag, str(prem_value))
    else:
        replace_tag_with_text(doc, '{REQ_PAY_5}', "__________")
        replace_tag_with_text(doc, '{REQ_COUNT_TOP5}', "__________")
        replace_tag_with_text(doc, '{REQ_PAY_WORDS_5}', "__________")
        for i in range(10, 101, 10):
            tag = '{PREM_TOP5_' + str(i) + '}'
            replace_tag_with_text(doc, tag, "__________")


def calculate_and_replace_tags_3(doc, request):
    req_count_top3 = request.POST.get('req_count_top3')
    if req_count_top3:
        req_count_top3 = int(req_count_top3)
        region_name = request.POST.get('region_name')
        search_engine = request.POST.get('search_engine')

        google_req_cost = 0
        if search_engine in ['GOOGLE', 'YANDEX_GOOGLE']:
            google_req_cost = int(request.POST.get('google_req_cost'))

        if region_name == 'Москва':
            req_pay = int(request.POST.get('req_count_top3_msk'))
        elif region_name == 'Санкт-Петербург':
            req_pay = int(request.POST.get('req_count_top3_spb'))
        else:
            req_pay = int(request.POST.get('req_count_top3_other'))

        req_pay += google_req_cost

        replace_tag_with_text(doc, '{REQ_PAY_3}', str(req_pay))
        replace_tag_with_text(doc, '{REQ_PAY_WORDS_3}', num2words(req_pay, lang='ru'))

        total_premium = req_pay * req_count_top3
        for i in range(10, 101, 10):
            percent = i / 100
            tag = '{PREM_TOP3_' + str(i) + '}'
            prem_value = int(total_premium * percent)
            replace_tag_with_text(doc, tag, str(prem_value))
    else:
        replace_tag_with_text(doc, '{REQ_PAY_3}', "__________")
        replace_tag_with_text(doc, '{REQ_COUNT_TOP3}', "__________")
        replace_tag_with_text(doc, '{REQ_PAY_WORDS_3}', "__________")
        for i in range(10, 101, 10):
            tag = '{PREM_TOP3_' + str(i) + '}'
            replace_tag_with_text(doc, tag, "__________")


def add_newline_before_text(doc, search_text):
    """
    Вставляет новую строку перед указанным текстом в документе.
    """
    for paragraph in doc.paragraphs:
        if search_text in paragraph.text:
            before_text, after_text = paragraph.text.split(search_text, 1)
            new_text = before_text.rstrip() + '\n' + search_text + after_text
            replace_paragraph_text_with_styles(paragraph, new_text)


def handle_conditional_sections(doc, edo):
    edo_text_1 = "(в том числе его получения с использованием системы электронного документооборота)" if edo == "YES" else ""
    edo_text_2 = (
            "\nЛибо посредством ЭДО: "
            "\n- ID Исполнителя в системе Тензор: 2BE894898d706174ab2aa3cdfc300550236"
            "\n- ID Заказчика: {CUSTOMER_ID} "
            "\n10.4. Стороны согласовали, что они вправе осуществлять документооборот в электронном виде по телекоммуникационным каналам связи с использованием усиленной квалификационной электронной подписи посредством системы электронного документооборота. " + "\n" +
            "10.4.1. В целях настоящего договора под электронным документом понимается документ, созданный в электронной форме без предварительного документирования на бумажном носителе, подписанный электронной подписью в порядке, установленном законодательством Российской Федерации. Стороны признают электронные документы, заверенные электронной подписью, при соблюдении требований Федерального закона от 06.04.2011 № 63-ФЗ 'Об электронной подписи' юридически эквивалентными документам на бумажных носителях, заверенным соответствующими подписями и оттиском печатей Сторон." + "\n" +
            "10.5. Все изменения и дополнения к договору оформляются в виде дополнений и приложений к договору, являющийся его неотъемлемой частью." + "\n" +
            "10.6. Договор составлен в двух подлинных экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из сторон, подписанных лично либо посредством ЭДО. ") \
        if edo == "YES" else ""
    not_edo_text = "на почту Исполнителя" if edo == "NO" else ""

    write_by_hand = (
        "\n10.4. Все изменения и дополнения к договору оформляются в виде дополнений и приложений к договору, являющийся его неотъемлемой частью."
        "10.5. Договор составлен в двух подлинных экземплярах, имеющих одинаковую юридическую силу, по одному для каждой из сторон.") if edo == "NO" else ""

    replacements = {
        '{EDO_1}': edo_text_1,
        '{EDO_2}': edo_text_2,
        '{NOT_EDO}': not_edo_text,
        '{WRITE_BY_HAND}': write_by_hand
    }

    for paragraph in doc.paragraphs:
        for tag, replacement in replacements.items():
            if tag in paragraph.text:
                paragraph_text = paragraph.text.replace(tag, replacement)
                replace_paragraph_text_with_styles(paragraph, paragraph_text)


def replace_tag_with_text(doc, tag, text=None):
    for paragraph in doc.paragraphs:
        if tag in paragraph.text:
            new_text = paragraph.text.replace(tag, text if text is not None else '').strip()
            if new_text:
                replace_paragraph_text_with_styles(paragraph, new_text)
            else:
                # Если параграф остался пустым, удаляем его
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None


def replace_underscores_with_signature(doc, placeholder_text, signature_path):
    def replace_in_paragraph(paragraph):
        runs = paragraph.runs
        full_text = ''.join(run.text for run in runs)
        if placeholder_text in full_text:
            split_text = full_text.split(placeholder_text)
            # Clear existing runs
            for run in runs:
                run.text = ""
            # Add new runs with the signature and styled text
            paragraph.add_run(split_text[0])
            paragraph.add_run().add_picture(signature_path, width=Inches(0.8))
            run = paragraph.add_run("Михайлов Д.С.")
            run.font.name = 'Calibri'
            run.font.size = Pt(9)
            if len(split_text) > 1:
                paragraph.add_run(split_text[1])

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)
    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            replace_in_paragraph(paragraph)
        footer = section.footer
        for paragraph in footer.paragraphs:
            replace_in_paragraph(paragraph)


def find_and_offset_director_text(doc):
    # Обход всех параграфов
    for paragraph in doc.paragraphs:
        if "________________{FIO_DIRECTOR}" in paragraph.text:
            paragraph.text = "\n\n\n" + paragraph.text

    # Обход всех секций (header и footer)
    for section in doc.sections:
        for header in section.header.paragraphs:
            if "________________{FIO_DIRECTOR}" in header.text:
                header.text = "\n\n\n" + header.text
        for footer in section.footer.paragraphs:
            if "________________{FIO_DIRECTOR}" in footer.text:
                footer.text = "\n\n\n" + footer.text

    # Обход всех таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "________________{FIO_DIRECTOR}" in paragraph.text:
                        paragraph.text = "\n\n\n" + paragraph.text


def process_contract(request):
    if request.method == 'POST':

        service_to_tag_mapping = {

            'optimization_headers': ('{HEAD_PAGES_1}', '- Оптимизация заголовков страниц;'),
            'optimization_headers2': (
                '{HEAD_PAGES_2}', '- Оптимизация заголовков страниц и техническое задание на их внедрение;'),
            'optimization_headers3': ('{HEAD_PAGES_3}', ''),

            'optimization_metatags': ('{METATAGS_1}', '- Оптимизация метатегов;'),
            'optimization_metatags2': (
                '{METATAGS_2}', '- Оптимизация метатегов и техническое задание на их внедрение;'),
            'optimization_metatags3': ('{METATAGS_3}', ''),

            'writing_optimization': ('{NEURO_1}', '- Написание текстов с помощью нейросетей и их оптимизация;'),
            'writing_optimization2': ('{NEURO_2}', '- Техническое задание на написание текстов и их оптимизация;'),
            'writing_optimization3': ('{NEURO_3}', ''),

            'site_structure_optimization': ('{STRUCTURE_1}', '- Оптимизация структуры сайта;'),
            'site_structure_optimization2': (
                '{STRUCTURE_2}', '- Оптимизация структуры сайта и техническое задание на внедрение;'),
            'site_structure_optimization3': ('{STRUCTURE_3}', ''),

            'technical_error_fixing': ('{FIX_1}', '- Устранение технических ошибок на сайте;'),
            'technical_error_fixing2': ('{FIX_2}', '- Техническое задание на устранение технических ошибок на сайте;'),
            'technical_error_fixing3': ('{FIX_3}', ''),

            'design_layouts': (
                '{TZ_1}',
                '- Техническое задание на создание дизайн-макетов отдельных блоков или страниц на сайте;'),

            'creating_pages': ('{CREATE_PAGES_1}', '- Создание страниц на сайте;'),
            'creating_pages2': ('{CREATE_PAGES_2}', '- Техническое задание на создание страниц на сайте;'),
            'creating_pages3': ('{CREATE_PAGES_3}', '')
        }

        contract_number = request.POST.get('contract_number')
        date_day = request.POST.get('date_day')
        site_name = request.POST.get('site_name')
        region_name = request.POST.get('region_name', 'СПб')
        custom_region_name = request.POST.get('custom_region_name', '').strip()

        if region_name == 'Другой' and custom_region_name:
            region_name = custom_region_name
        search_engine = request.POST.get('search_engine', 'YANDEX')

        price_count_digit = request.POST.get('price_count_digit')
        price_count_word = num2words(price_count_digit, lang='ru')

        date_month = request.POST.get('date_month')
        date_year = request.POST.get('date_year')

        organization_name = request.POST.get('organization_name')

        if organization_name.startswith('Индивидуальный предприниматель'):
            organization_name += ', именуемый'
        elif organization_name.startswith('ООО'):
            organization_name += ', именуемое'

        red_organization_name = request.POST.get('red_organization_name')
        guarantee = request.POST.get('guarantee')

        reason = request.POST.get('reason')
        person_name = request.POST.get('person_name')
        director_name = request.POST.get('director_name')
        email = request.POST.get('email')
        inn = request.POST.get('inn')
        customer_id = request.POST.get('customer_id')

        ogrn = request.POST.get('ogrn')
        registration_address = request.POST.get('registration_address')
        checking_account = request.POST.get('checking_account')
        correspondent_account = request.POST.get('correspondent_account')
        bank_name = request.POST.get('bank_name')
        bic = request.POST.get('bic')
        edo = request.POST.get('edo')
        req_count_top10 = request.POST.get('req_count_top10')
        req_count_top5 = request.POST.get('req_count_top5')
        req_count_top3 = request.POST.get('req_count_top3')
        topvisor = request.POST.get('topvisor')
        search_engine = request.POST.get('search_engine')
        support_options = request.POST.getlist('support[]')
        platform_choice = request.POST.get('platform', None)
        selected_services = request.POST.getlist('services[]')
        choose_executor = request.POST.get('choose_executor')

        template_filename = 'Договор Позиции метки.docx'
        template_path = os.path.join(os.path.dirname(__file__), '../dogovora', template_filename)
        doc = Document(template_path)
        signature_image_path = os.path.join(os.path.dirname(__file__), '../dogovora/podpis.jpg')

        if choose_executor == 'ИП Михайлов Дмитрий Сергеевич':
            replace_tag_with_text(doc, '{PREDMET_DOGOVORA1}', 'услуги по адаптации и оптимизации web-страниц')
            executor_name_replacement = ('Индивидуальный предприниматель Михайлов Дмитрий Сергеевич, именуемый в '
                                         'дальнейшем «Исполнитель», в лице генерального директора Михайлова Дмитрия '
                                         'Сергеевича, действующего '
                                         'на основании Свидетельства ОГРНИП 320784700136130')
            replacements_executor = {
                '{CHOOSE_EXECUTOR_NAME}': 'Индивидуальный предприниматель Михайлов Дмитрий Сергеевич',
                '{CHOOSE_EXECUTOR_INN}': '780256693210',
                '{CHOOSE_EXECUTOR_OGRNIP}': 'ОГРНИП: 320784700136130',
                '{CHOOSE_EXECUTOR_ADRESS}': '194295, Россия, г. Санкт-Петербург, пр-кт Северный, д. 24, корпус 1, '
                                            'кв. 33',
                '{CHOOSE_EXECUTOR_CHECKING_ACC}': '40802810201500152101',
                '{CHOOSE_EXECUTOR_KOR_ACC}': '30101810745374525104',
                '{CHOOSE_EXECUTOR_BANK}': 'ООО "Банк Точка"',
                '{CHOOSE_EXECUTOR_BIK}': '044525104',
                '{CHOOSE_EXECUTOR_EMAIL}': 'info@mihaylov.digital'
            }
        elif choose_executor == 'ООО «МД»':
            replace_tag_with_text(doc, '{PREDMET_DOGOVORA1}', 'рекламные услуги по поисковой оптимизации')
            executor_name_replacement = ('Общество с ограниченной ответственностью "Михайлов Диджитал", именуемый в '
                                         'дальнейшем «Исполнитель», в лице генерального директора Михайлова Дмитрия '
                                         'Сергеевича, действующего '
                                         'на основании Устава')
            replacements_executor = {
                '{CHOOSE_EXECUTOR_NAME}': 'Общество с ограниченной ответственностью "Михайлов Диджитал"',
                '{CHOOSE_EXECUTOR_INN}': '7810962062',
                '{CHOOSE_EXECUTOR_OGRNIP}': 'ОГРН: 1247800061464',
                '{CHOOSE_EXECUTOR_ADRESS}': '196142, Россия, г. Санкт-Петербург, ул. Пулковская, д. 2, корпус 1, '
                                            'литера А, оф 25, помещ. 66-Н',
                '{CHOOSE_EXECUTOR_CHECKING_ACC}': '40702810320000118082',
                '{CHOOSE_EXECUTOR_KOR_ACC}': '30101810745374525104',
                '{CHOOSE_EXECUTOR_BANK}': 'ООО "Банк Точка"',
                '{CHOOSE_EXECUTOR_BIK}': '044525104',
                '{CHOOSE_EXECUTOR_EMAIL}': 'info@mihaylov.digital'
            }
        replace_tag_with_text(doc, '{CHOOSE_EXECUTOR_NAME}', executor_name_replacement)

        for paragraph in doc.paragraphs:
            for key, value in replacements_executor.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, value)
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements_executor.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)
                                for run in paragraph.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(9)

        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                for key, value in replacements_executor.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, value)
                        for run in paragraph.runs:
                            run.font.name = 'Calibri'
                            run.font.size = Pt(9)

        calculate_and_replace_tags_10(doc, request)
        calculate_and_replace_tags_5(doc, request)
        calculate_and_replace_tags_3(doc, request)

        for paragraph in doc.paragraphs:
            if '{YANDEX}' in paragraph.text or '{GOOGLE}' in paragraph.text or '{YANDEX_GOOGLE}' in paragraph.text:
                if search_engine == 'YANDEX':
                    paragraph_text = paragraph.text.replace('{YANDEX}', 'Яндекс').replace('{GOOGLE}', '').replace(
                        '{YANDEX_GOOGLE}', '')
                elif search_engine == 'GOOGLE':
                    paragraph_text = paragraph.text.replace('{YANDEX}', '').replace('{GOOGLE}', 'Google').replace(
                        '{YANDEX_GOOGLE}', '')
                else:
                    paragraph_text = paragraph.text.replace('{YANDEX}', '').replace('{GOOGLE}', '').replace(
                        '{YANDEX_GOOGLE}', 'Яндекс и Google')
                replace_paragraph_text_with_styles(paragraph, paragraph_text)

        for paragraph in doc.paragraphs:
            if '{REGION_NAME}' in paragraph.text:
                paragraph_text = paragraph.text.replace('{REGION_NAME}', region_name)
                replace_paragraph_text_with_styles(paragraph, paragraph_text)

        if 'support_site' in support_options:
            support_text = '- Поддержка технического состояния сайта;'
            replace_tag_with_text(doc, '{SITE_SUPPORT}', support_text)
        else:
            replace_tag_with_text(doc, '{SITE_SUPPORT}')

        used_tags = set()

        for service_key in selected_services:
            if service_key in service_to_tag_mapping:
                tag, text = service_to_tag_mapping[service_key]
                replace_tag_with_text(doc, tag, text)
                used_tags.add(tag)

        all_tags = set(service_to_tag_mapping.values())
        unused_tags = {tag for tag, _ in all_tags if tag not in used_tags}
        for tag in unused_tags:
            replace_tag_with_text(doc, tag)

        handle_conditional_sections(doc, edo)
        handle_additional_work_sections(doc, selected_services, platform_choice)

        if guarantee == 'with_guarantee':
            add_guarantee_section(doc)

        if search_engine == 'YANDEX':
            # Изменение пункта 1.4.1
            remove_bullet_point(doc, '- Анализ ссылочного профиля сайта;')

            # Изменение пункта 1.4.6

        footer = doc.sections[0].footer
        for paragraph in footer.paragraphs:
            paragraph.alignment = 1
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
        paragraph = footer.paragraphs[0]
        paragraph.add_run("________________" + director_name + "").font.size = Pt(12)
        paragraph.add_run("                                                                           ").font.size = Pt(
            12)

        if edo == "YES":
            paragraph.add_run("_______________Михайлов Д.С.").font.size = Pt(12)
        else:
            replace_underscores_with_signature(doc, "_______________Михайлов Д.С.", signature_image_path)
            find_and_offset_director_text(doc)
            run = paragraph.add_run()
            run.add_picture(signature_image_path, width=Inches(0.8))  # Настройте ширину по необходимости
            paragraph.add_run("Михайлов Д.С.").font.size = Pt(12)

        add_newline_before_text(doc, "10.5. Договор составлен")
        replacements = {
            '{DOGOVOR_NUMBER}': contract_number,
            '{DAY}': date_day,
            '{MONTH}': date_month,
            '{YEAR}': date_year,
            '{CUSTOMER_ORGANIZATION}': organization_name,
            '{CUSTOMER_FIO}': person_name,
            '{DOGOVOR_OSNOVANIE}': reason,

            '{SITE_NAME}': site_name,
            '{REGION_NAME}': region_name,
            '{PRICE_COUNT}': price_count_digit,
            '{PRICE_COUNT_IN_WORDS}': price_count_word,

            '{FIO_DIRECTOR}': director_name,
            '{CUSTOMER_EMAIL}': email,
            '{CUSTOMER_NAME}': organization_name,
            '{RED_CUSTOMER_ORGANIZATION}': red_organization_name,
            '{INN}': inn,
            '{CUSTOMER_ID}': customer_id,

            '{OGRN}': ogrn,
            '{REGISTRATION_ADDRESS}': registration_address,
            '{PAYMENT_ACCOUNT}': checking_account,
            '{CORRESPONDENT}': correspondent_account,
            '{BANK_NAME}': bank_name,
            '{BIK}': bic,
            '{REQ_COUNT_TOP10}': req_count_top10,
            '{REQ_COUNT_TOP5}': req_count_top5,
            '{REQ_COUNT_TOP3}': req_count_top3,
            '{TOPVISOR}': topvisor,
            '{CHOOSE_EXECUTOR}': choose_executor,

        }

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                if key in paragraph.text:
                    paragraph.text = paragraph.text.replace(key, str(value))
                    for run in paragraph.runs:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(9)
                        if key == '{DOGOVOR_NUMBER}' or key == '{CUSTOMER_NAME}':
                            run.bold = True

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in paragraph.text:
                                paragraph.text = paragraph.text.replace(key, value)
                                for run in paragraph.runs:
                                    run.font.name = 'Calibri'
                                    run.font.size = Pt(9)
                                    if run.text.strip() == value:
                                        run.bold = True

        for section in doc.sections:
            footer = section.footer
            for paragraph in footer.paragraphs:
                paragraph.alignment = 1
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9)

        text = organization_name
        make_text_bold_in_doc(doc, text)
        make_text_bold_in_doc(doc, 'Индивидуальный предприниматель Михайлов Дмитрий Сергеевич')
        make_text_bold_in_doc(doc, 'Общество с ограниченной ответственностью "Михайлов Диджитал"')

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="processed_contract.docx"'
        doc.save(response)
        return response
    else:
        return render(request, 'dogovor_position_form.html')
