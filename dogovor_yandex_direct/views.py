import os
from django.http import HttpResponse
from django.shortcuts import render
from docx import Document
from docx.shared import Pt, Inches
from num2words import num2words


def add_styled_run(paragraph, text, bold=False, size=Pt(9)):
    run = paragraph.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = size
    run.bold = bold
    return run


def replace_tags_in_doc(doc, replacements):
    # Абзацы в основном теле документа
    for paragraph in doc.paragraphs:
        replace_paragraph_text(paragraph, replacements)

    # Абзацы внутри таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_paragraph_text(paragraph, replacements)

    # Абзацы в хедерах и футерах
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            replace_paragraph_text(paragraph, replacements)
        for paragraph in section.footer.paragraphs:
            replace_paragraph_text(paragraph, replacements)


def replace_paragraph_text(paragraph, replacements):
    for key, value in replacements.items():
        if key in paragraph.text:
            paragraph.text = paragraph.text.replace(key, str(value))
            for run in paragraph.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(9)


def process_contract(request):
    if request.method == 'POST':
        contract_number = request.POST.get('contract_number')
        date_day = request.POST.get('date_day')
        date_month = request.POST.get('date_month')
        date_year = request.POST.get('date_year')
        organization_name = request.POST.get('organization_name')
        reason = request.POST.get('reason')
        person_name = request.POST.get('person_name')

        choose_executor = request.POST.get('choose_executor')
        site_name = request.POST.get('site_name')

        price_count = request.POST.get('price_count')
        price_count_in_words = num2words(price_count, lang='ru')

        customer_email = request.POST.get('email')
        customer_id = request.POST.get('customer_id')

        red_organization_name = request.POST.get('red_organization_name')

        director_name = request.POST.get('director_name')
        inn = request.POST.get('inn')
        ogrn = request.POST.get('ogrn')
        registration_address = request.POST.get('registration_address')
        checking_account = request.POST.get('checking_account')
        correspondent_account = request.POST.get('correspondent_account')
        bank_name = request.POST.get('bank_name')
        bic = request.POST.get('bic')
        edo = request.POST.get('edo')

        if choose_executor == 'ИП Михайлов Дмитрий Сергеевич':
            predmet_1 = 'адаптации и оптимизации web-страниц'
            predmet_2 = 'адаптации и оптимизации web-страниц'
            predmet_3 = 'адаптации и оптимизации web-страниц'

            replacements_executor = {
                '{CHOOSE_EXECUTOR_NAME}': 'Индивидуальный предприниматель Михайлов Дмитрий Сергеевич',
                '{CHOOSE_EXECUTOR_INN}': '780256693210',
                '{CHOOSE_EXECUTOR_OGRNIP}': 'ОГРНИП: 320784700136130',
                '{CHOOSE_EXECUTOR_ADRESS}': '194295, Россия, г. Санкт-Петербург, пр-кт Северный, д. 24, корпус 1, кв. 33',
                '{CHOOSE_EXECUTOR_CHECKING_ACC}': '40802810201500152101',
                '{CHOOSE_EXECUTOR_KOR_ACC}': '30101810745374525104',
                '{CHOOSE_EXECUTOR_BANK}': 'ООО "Банк Точка"',
                '{CHOOSE_EXECUTOR_BIK}': '044525104',
                '{CHOOSE_EXECUTOR_EMAIL}': 'info@mihaylov.digital'
            }

        else:
            predmet_1 = 'поисковой оптимизации'
            predmet_2 = 'поисковой оптимизации'
            predmet_3 = 'поисковой оптимизации'

            replacements_executor = {
                '{CHOOSE_EXECUTOR_NAME}': 'Общество с ограниченной ответственностью "Михайлов Диджитал"',
                '{CHOOSE_EXECUTOR_INN}': '7810962062',
                '{CHOOSE_EXECUTOR_OGRNIP}': 'ОГРН: 1247800061464',
                '{CHOOSE_EXECUTOR_ADRESS}': '196142, Россия, г. Санкт-Петербург, ул. Пулковская, д. 2, корпус 1, литера А, оф 25, помещ. 66-Н',
                '{CHOOSE_EXECUTOR_CHECKING_ACC}': '40702810320000118082',
                '{CHOOSE_EXECUTOR_KOR_ACC}': '30101810745374525104',
                '{CHOOSE_EXECUTOR_BANK}': 'ООО "Банк Точка"',
                '{CHOOSE_EXECUTOR_BIK}': '044525104',
                '{CHOOSE_EXECUTOR_EMAIL}': 'info@mihaylov.digital'
            }

            # Добавляем реквизиты исполнителя в общий список замен

        template_path = os.path.join(os.path.dirname(__file__), '../dogovora/Договор Яндекс директ метки.docx')
        doc = Document(template_path)

        # Обработка вставки по ЭДО в пункт 10.3
        for paragraph in doc.paragraphs:
            if '{EDO}' in paragraph.text:
                if edo == 'NO':
                    lines = paragraph.text.split('\n')
                    cleaned_lines = []
                    for line in lines:
                        # Убираем строки с ЭДО
                        if not any(x in line for x in ['Либо посредством ЭДО', 'ID Исполнителя', 'ID Заказчика']):
                            cleaned_lines.append(line.replace('{EDO}', ''))
                    paragraph.text = '\n'.join(cleaned_lines)
                else:
                    paragraph.text = paragraph.text.replace('{EDO}', '')

                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9)

        replacements = {
            '{DOGOVOR_NUMBER}': contract_number,
            '{DAY}': date_day,
            '{MONTH}': date_month,
            '{YEAR}': date_year,
            '{CUSTOMER_ORGANIZATION}': organization_name,
            '{DOGOVOR_OSNOVANIE}': reason,
            '{CUSTOMER_FIO}': person_name,
            '{PREDMET_DOGOVORA1}': predmet_1,
            '{PREDMET_DOGOVORA2}': predmet_2,
            '{PREDMET_DOGOVORA3}': predmet_3,
            '{SITE_NAME}': site_name,
            '{PRICE_COUNT}': price_count,
            '{PRICE_COUNT_IN_WORDS}': price_count_in_words,

            '{CUSTOMER_EMAIL}': customer_email,
            '{CUSTOMER_ID}': customer_id,

            '{RED_CUSTOMER_ORGANIZATION}': red_organization_name,

            '{FIO_DIRECTOR}': director_name,
            '{CUSTOMER_NAME}': organization_name,
            '{INN}': inn,
            '{OGRN}': ogrn,
            '{REGISTRATION_ADDRESS}': registration_address,
            '{PAYMENT_ACCOUNT}': checking_account,
            '{CORRESPONDENT}': correspondent_account,
            '{BANK_NAME}': bank_name,
            '{BIK}': bic,
            '{CHOOSE_EXECUTOR}': choose_executor,
        }
        replacements.update(replacements_executor)

        # Замена во всех параграфах
        replace_tags_in_doc(doc, replacements)

        signature_image_path = os.path.join(os.path.dirname(__file__), '../dogovora/podpis.jpg')

        # Обработка колонтитула для подписи
        footer = doc.sections[0].footer
        footer.paragraphs.clear()  # очистим, чтобы не было наложений
        paragraph = footer.add_paragraph()
        paragraph.alignment = 1  # Центровка

        # Левая сторона – Заказчик
        run_left = paragraph.add_run("________________" + (director_name or ''))
        run_left.font.name = 'Calibri'
        run_left.font.size = Pt(9)

        # Пробелы между
        paragraph.add_run(" " * 30)  # достаточно, чтобы разделить подписи визуально

        # Правая сторона – Исполнитель
        if edo == "YES":
            run_exec = paragraph.add_run("_______________Михайлов Д.С.")
            run_exec.font.name = 'Calibri'
            run_exec.font.size = Pt(9)
        else:
            run_pic = paragraph.add_run()
            run_pic.add_picture(signature_image_path, width=Inches(0.8))
            run_text = paragraph.add_run("Михайлов Д.С.")
            run_text.font.name = 'Calibri'
            run_text.font.size = Pt(9)

        # Отдаём как docx-файл
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename="contract_yandex_direct.docx"'
        doc.save(response)
        return response

    # GET-запрос
    return render(request, '../templates/dogovor_yandex_direct.html')
